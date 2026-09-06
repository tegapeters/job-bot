"""
Job Pal — Streamlit UI v2 (Techturi branded)
Run: streamlit run ui_v2.py
"""
import io
import re
from datetime import datetime, timezone
import streamlit as st
import pandas as pd
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from tracker import (
    get_all_applications, get_review_queue,
    update_status, get_seen_ids, log_event, get_event_counts,
    log_experiment_run, get_recent_runs,
    rank_queue_with_personalization, get_source_health,
    clear_queue, upsert_events, get_events, update_event_status, delete_past_events, delete_all_events,
    _scope_id,
)
from sessions import save_session, load_session, clear_session, new_uid, save_chat_history, load_chat_history, clear_chat_history, save_gmail_opt_in
from auth import render_auth_wall, restore_user_session, get_user_id, get_user_email, sign_out
from config import REVIEW_MIN_SCORE

# ── Page config ────────────────────────────────────────────────────
st.set_page_config(
    page_title="Job Pal · techturi",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="auto",
)

# ── Brand styles ───────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,300..900;1,9..144,300..900&family=JetBrains+Mono:wght@400;500;600&display=swap');
  /* ── Base ── */
  html, body, [data-testid="stApp"] { background: #0A0A0B; }
  [data-testid="stSidebar"] {
    background: #0f0f10;
    border-right: 1px solid #1f1f22;
  }
  [data-testid="stSidebar"] > div { padding-top: 0 !important; }

  /* ── Logo ── */
  .tt-logo {
    font-family: 'JetBrains Mono', monospace;
    font-size: 20px;
    font-weight: 500;
    color: #F5F4EE;
    letter-spacing: -0.01em;
    padding: 28px 20px 0;
  }
  .tt-logo .bracket, .tt-logo .dot { color: #D4FF3A; }
  .tt-product {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.25em;
    color: #4A4A45;
    text-transform: uppercase;
    padding: 6px 20px 20px;
    border-bottom: 1px solid #1f1f22;
    margin-bottom: 12px;
  }

  /* ── Page header ── */
  .page-header {
    display: flex;
    align-items: center;
    gap: 16px;
    padding-bottom: 20px;
    border-bottom: 1px solid #1f1f22;
    margin-bottom: 28px;
  }
  .page-header .eyebrow {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.25em;
    color: #D4FF3A;
    text-transform: uppercase;
    margin-bottom: 4px;
  }
  .page-header h1 {
    font-family: 'Fraunces', serif !important;
    font-size: 40px !important;
    font-weight: 300 !important;
    color: #F5F4EE !important;
    letter-spacing: -0.03em !important;
    line-height: 1 !important;
    margin: 0 !important;
    padding: 0 !important;
  }
  .page-header h1 em {
    font-style: italic;
    color: #D4FF3A;
  }

  /* ── Metric cards ── */
  .metric-card {
    background: #131315;
    border: 1px solid #1f1f22;
    border-radius: 8px;
    padding: 20px 24px;
    margin-bottom: 12px;
  }
  .metric-card .label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #4A4A45;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    margin-bottom: 8px;
  }
  .metric-card .value {
    font-family: 'Fraunces', serif;
    font-size: 36px;
    font-weight: 400;
    color: #F5F4EE;
    line-height: 1;
  }
  .metric-card .sub {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #4A4A45;
    margin-top: 6px;
  }
  .metric-card.accent { border-left: 3px solid #D4FF3A; }

  /* ── Score / status badges ── */
  .score-high { color: #D4FF3A; font-weight: 600; font-family: 'JetBrains Mono', monospace; }
  .score-mid  { color: #f5c518; font-weight: 600; font-family: 'JetBrains Mono', monospace; }
  .score-low  { color: #ff6b6b; font-weight: 600; font-family: 'JetBrains Mono', monospace; }
  .tag {
    display: inline-block;
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    padding: 3px 10px;
    border-radius: 2px;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    margin-right: 4px;
  }
  .tag-new                { background: #1a1f0a; color: #D4FF3A; border: 1px solid #2a3510; }
  .tag-applied            { background: #0a1a1f; color: #3ad4ff; border: 1px solid #102a35; }
  .tag-interview          { background: #1a0f1a; color: #d43aff; border: 1px solid #2a1035; }
  .tag-rejected           { background: #1f0a0a; color: #ff3a3a; border: 1px solid #350f0f; }
  .tag-skipped            { background: #1a1a1a; color: #666; border: 1px solid #333; }
  .tag-application_closed { background: #1a1208; color: #ff9a3a; border: 1px solid #352010; }
  .tag-no_response        { background: #181818; color: #888; border: 1px solid #2a2a2a; }

  /* ── Cover letter block ── */
  .cover-letter {
    background: #131315;
    border: 1px solid #1f1f22;
    border-left: 3px solid #D4FF3A;
    border-radius: 4px;
    padding: 20px 24px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 13px;
    color: #c8c8c0;
    white-space: pre-wrap;
    line-height: 1.8;
  }

  /* ── Section divider label ── */
  .section-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.2em;
    color: #4A4A45;
    text-transform: uppercase;
    margin-bottom: 16px;
    padding-bottom: 10px;
    border-bottom: 1px solid #1f1f22;
  }

  /* ── Pipeline status box ── */
  .pipeline-card {
    background: #131315;
    border: 1px solid #1f1f22;
    border-radius: 8px;
    padding: 24px;
  }
  .pipeline-card h4 {
    font-family: 'Fraunces', serif !important;
    font-size: 22px !important;
    font-weight: 400 !important;
    color: #F5F4EE !important;
    margin: 0 0 8px 0 !important;
  }
  .pipeline-card p {
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    color: #8B8B85;
    line-height: 1.7;
    margin: 0 !important;
  }

  /* ── All buttons — base reset ── */
  div.stButton > button,
  div[data-testid="stButton"] > button,
  div[data-testid="stDownloadButton"] > button,
  div[data-testid="stFormSubmitButton"] > button,
  [data-testid="baseButton-primary"],
  [data-testid="baseButton-secondary"] {
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 12px !important;
    letter-spacing: 0.05em !important;
    border-radius: 3px !important;
    transition: background 0.15s, border-color 0.15s !important;
  }

  /* ── Primary buttons (white) ── */
  div.stButton > button[kind="primary"],
  div[data-testid="stButton"] > button[kind="primary"],
  div[data-testid="stFormSubmitButton"] > button[kind="primary"],
  [data-testid="baseButton-primary"] {
    background: #F5F4EE !important;
    color: #0A0A0B !important;
    border: 1px solid #F5F4EE !important;
    font-weight: 600 !important;
  }
  div.stButton > button[kind="primary"]:hover,
  div[data-testid="stButton"] > button[kind="primary"]:hover,
  [data-testid="baseButton-primary"]:hover {
    background: #ffffff !important;
    border-color: #ffffff !important;
  }

  /* ── Secondary / default buttons (dark grey) ── */
  div.stButton > button[kind="secondary"],
  div.stButton > button:not([kind="primary"]),
  div[data-testid="stButton"] > button[kind="secondary"],
  div[data-testid="stDownloadButton"] > button,
  [data-testid="baseButton-secondary"] {
    background: #1f1f22 !important;
    color: #c8c8c0 !important;
    border: 1px solid #2e2e32 !important;
    font-weight: 500 !important;
  }
  div.stButton > button[kind="secondary"]:hover,
  div.stButton > button:not([kind="primary"]):hover,
  div[data-testid="stButton"] > button[kind="secondary"]:hover,
  div[data-testid="stDownloadButton"] > button:hover,
  [data-testid="baseButton-secondary"]:hover {
    background: #2a2a2e !important;
    border-color: #3e3e44 !important;
    color: #F5F4EE !important;
  }

  /* ── Multiselect tags ── */
  /* Target all variants including those with Streamlit-injected inline style= */
  span[data-baseweb="tag"],
  span[data-baseweb="tag"][style],
  [data-testid="stMultiSelect"] span[data-baseweb="tag"],
  [data-testid="stMultiSelect"] span[data-baseweb="tag"][style],
  [data-baseweb="select"] span[data-baseweb="tag"],
  div[data-baseweb="select"] span[data-baseweb="tag"] {
    background: #2a2a2e !important;
    background-color: #2a2a2e !important;
    color: #c8c8c0 !important;
    border: 1px solid #3e3e44 !important;
    border-radius: 4px !important;
  }
  span[data-baseweb="tag"] span,
  span[data-baseweb="tag"] > span,
  span[data-baseweb="tag"][style] span,
  [data-testid="stMultiSelect"] span[data-baseweb="tag"] span,
  [data-testid="stMultiSelect"] span[data-baseweb="tag"] > span {
    color: #c8c8c0 !important;
    background: transparent !important;
    background-color: transparent !important;
  }
  /* X button inside tags */
  span[data-baseweb="tag"] [role="presentation"],
  span[data-baseweb="tag"] button,
  span[data-baseweb="tag"] svg {
    color: #888 !important;
    fill: #888 !important;
    background: transparent !important;
    background-color: transparent !important;
  }

  /* ── Slider thumb + track ── */
  [data-testid="stSlider"] [role="slider"] {
    background: #c8c8c0 !important;
    border-color: #c8c8c0 !important;
  }
  [data-testid="stSlider"] [data-testid="stSliderThumbValue"] {
    color: #c8c8c0 !important;
  }

  /* ── Footer ── */
  .tt-footer {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #4A4A45;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    padding-top: 8px;
    border-top: 1px solid #1f1f22;
    margin-top: 4px;
  }

  /* ══════════════════════════════════════════════════════════════
     MOBILE  (≤ 768px)
     ══════════════════════════════════════════════════════════════ */
  @media (max-width: 768px) {

    /* ── Main content padding ── */
    .main .block-container {
      padding-left: 16px !important;
      padding-right: 16px !important;
      padding-top: 16px !important;
    }

    /* ── Stack all Streamlit columns ── */
    [data-testid="stHorizontalBlock"] {
      flex-wrap: wrap !important;
      gap: 0 !important;
    }
    [data-testid="column"] {
      min-width: 100% !important;
      width: 100% !important;
      flex: 1 1 100% !important;
    }

    /* ── Page header ── */
    .page-header { margin-bottom: 20px; padding-bottom: 16px; }
    .page-header h1 {
      font-size: 28px !important;
      letter-spacing: -0.02em !important;
    }

    /* ── Auth / landing hero ── */
    div[style*="max-width:440px"],
    div[style*="max-width: 440px"] {
      margin-top: 32px !important;
    }
    div[style*="font-size:42px"],
    div[style*="font-size: 42px"] {
      font-size: 30px !important;
    }

    /* ── Metric cards: smaller value, tighter padding ── */
    .metric-card { padding: 14px 16px; }
    .metric-card .value { font-size: 26px; }

    /* ── Buttons: bigger tap targets ── */
    div.stButton > button,
    div[data-testid="stButton"] > button,
    div[data-testid="stDownloadButton"] > button,
    div[data-testid="stFormSubmitButton"] > button {
      min-height: 48px !important;
      font-size: 13px !important;
    }

    /* ── Cover letter: smaller mono text ── */
    .cover-letter {
      font-size: 11px !important;
      padding: 14px 16px !important;
      line-height: 1.7 !important;
    }

    /* ── Pipeline card ── */
    .pipeline-card { padding: 16px; }
    .pipeline-card h4 { font-size: 18px !important; }

    /* ── Section label spacing ── */
    .section-label { margin-bottom: 12px; }

    /* ── Tabs: smaller text ── */
    [data-testid="stTabs"] button {
      font-size: 11px !important;
      padding: 6px 10px !important;
    }

    /* ── Expander header ── */
    [data-testid="stExpander"] summary {
      font-size: 13px !important;
      padding: 10px 14px !important;
    }

    /* ── Sidebar logo on mobile ── */
    .tt-logo { font-size: 17px !important; padding: 20px 16px 0 !important; }
    .tt-product { padding: 4px 16px 16px !important; }

    /* ── Sliders: wider touch area ── */
    [data-testid="stSlider"] [role="slider"] {
      width: 22px !important;
      height: 22px !important;
    }

    /* ── Dataframes: allow horizontal scroll ── */
    [data-testid="stDataFrame"] {
      overflow-x: auto !important;
    }
    [data-testid="stDataFrame"] > div {
      min-width: 480px;
    }

    /* ── Sidebar toggle: larger tap target + accent colour ── */
    [data-testid="collapsedControl"] {
      width: 44px !important;
      height: 44px !important;
      background: #D4FF3A !important;
      border-radius: 8px !important;
      display: flex !important;
      align-items: center !important;
      justify-content: center !important;
    }
    [data-testid="collapsedControl"] svg {
      color: #0A0A08 !important;
      fill: #0A0A08 !important;
      width: 20px !important;
      height: 20px !important;
    }
  }

  /* ── Landscape phones (height ≤ 500px) ── */
  @media (max-height: 500px) and (max-width: 900px) {
    div[style*="max-width:440px"],
    div[style*="max-width: 440px"] {
      margin-top: 8px !important;
    }
    div[style*="font-size:42px"],
    div[style*="font-size: 42px"] {
      font-size: 20px !important;
      margin-bottom: 8px !important;
      line-height: 1.2 !important;
    }
    div[style*="margin-bottom:10px"] {
      margin-bottom: 4px !important;
    }
  }

  /* ── Hide Streamlit health status badge ── */
  [data-testid="stStatusWidget"] { display: none !important; }

  /* ── Swipe card (Review Queue card mode) ── */
  .swipe-card {
    background: #131315;
    border: 1px solid #1f1f22;
    border-radius: 12px;
    padding: 28px 28px 24px;
    margin: 0 auto 24px;
    max-width: 600px;
    position: relative;
    touch-action: pan-y;
    user-select: none;
    -webkit-user-select: none;
  }
  .sc-progress-bar {
    height: 3px;
    background: #1f1f22;
    border-radius: 2px;
    margin-bottom: 20px;
    max-width: 600px;
    margin-left: auto;
    margin-right: auto;
  }
  .sc-progress-fill {
    height: 100%;
    background: #D4FF3A;
    border-radius: 2px;
    transition: width 0.3s ease;
  }
  .sc-counter {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #4A4A45;
    letter-spacing: 0.15em;
    text-align: center;
    margin-bottom: 20px;
  }
  .sc-score-badge {
    font-family: 'Fraunces', serif;
    font-size: 52px;
    font-weight: 300;
    line-height: 1;
    margin-bottom: 4px;
  }
  .sc-score-badge.high { color: #D4FF3A; }
  .sc-score-badge.mid  { color: #f5c518; }
  .sc-score-badge.low  { color: #ff6b6b; }
  .sc-score-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #4A4A45;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    margin-bottom: 20px;
  }
  .sc-title {
    font-family: 'Fraunces', serif;
    font-size: 26px;
    font-weight: 300;
    color: #F5F4EE;
    margin: 0 0 6px 0;
    line-height: 1.2;
  }
  .sc-company {
    font-family: 'JetBrains Mono', monospace;
    font-size: 13px;
    color: #8B8B85;
    margin-bottom: 16px;
  }
  .sc-tags { display: flex; gap: 8px; flex-wrap: wrap; margin-bottom: 18px; }
  .sc-tag {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.1em;
    padding: 4px 10px;
    border-radius: 3px;
    text-transform: uppercase;
  }
  .sc-tag.remote { background: #1a1f0a; color: #D4FF3A; border: 1px solid #2a3510; }
  .sc-tag.hybrid { background: #0a1a1f; color: #3ad4ff; border: 1px solid #102a35; }
  .sc-tag.onsite { background: #1a1a1a; color: #8B8B85; border: 1px solid #2a2a2a; }
  .sc-tag.salary { background: #131315; color: #8B8B85; border: 1px solid #2a2a2a; }
  .sc-tag.days   { background: #131315; color: #8B8B85; border: 1px solid #2a2a2a; }
  .sc-tag.days.warn { color: #f5c518; }
  .sc-tag.days.stale { color: #ff6b6b; }
  .sc-reason {
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    color: #c8c8c0;
    line-height: 1.7;
    margin-bottom: 14px;
    padding: 12px 14px;
    background: #0d0d0f;
    border-radius: 6px;
    border-left: 2px solid #D4FF3A;
  }
  .sc-hint {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #8B8B85;
    margin-bottom: 14px;
  }
  .sc-link {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #D4FF3A;
    text-decoration: none;
    letter-spacing: 0.08em;
  }
  .sc-link:hover { text-decoration: underline; }
  .sc-gesture-hint {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #2a2a2e;
    text-align: center;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    margin-top: 14px;
  }
  .sc-swipe-indicator {
    position: absolute;
    top: 20px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    font-weight: 600;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    padding: 5px 12px;
    border-radius: 4px;
    opacity: 0;
    pointer-events: none;
    transition: opacity 0.05s;
  }
  .sc-swipe-indicator.skip-ind  { left: 20px;  color: #ff6b6b; border: 2px solid #ff6b6b; }
  .sc-swipe-indicator.apply-ind { right: 20px; color: #D4FF3A; border: 2px solid #D4FF3A; }
  @media (max-width: 768px) {
    .swipe-card { padding: 20px 18px 18px; border-radius: 10px; }
    .sc-title { font-size: 22px; }
    .sc-score-badge { font-size: 42px; }
  }

  /* ── Funnel row (Dashboard) ── */
  .funnel-row {
    display: flex;
    align-items: center;
    gap: 0;
    background: #131315;
    border: 1px solid #1f1f22;
    border-radius: 8px;
    padding: 24px 28px;
    margin-bottom: 28px;
    flex-wrap: wrap;
  }
  .funnel-step { text-align: center; flex: 1; min-width: 80px; }
  .funnel-step .fn { font-family: 'Fraunces', serif; font-size: 32px; font-weight: 400; color: #F5F4EE; line-height: 1; }
  .funnel-step .fn.accent { color: #D4FF3A; }
  .funnel-step .fl { font-family: 'JetBrains Mono', monospace; font-size: 10px; color: #4A4A45; letter-spacing: 0.15em; text-transform: uppercase; margin-top: 6px; }
  .funnel-arrow { font-family: 'JetBrains Mono', monospace; font-size: 16px; color: #2e2e32; padding: 0 12px; flex-shrink: 0; }

  /* ── Pre-flight checklist (Run Pipeline) ── */
  .preflight {
    background: #131315;
    border: 1px solid #1f1f22;
    border-radius: 8px;
    padding: 20px 24px;
    margin-bottom: 20px;
  }
  .preflight .pf-row {
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    color: #8B8B85;
    padding: 4px 0;
    display: flex;
    align-items: center;
    gap: 10px;
  }
  .preflight .pf-row .ok  { color: #D4FF3A; }
  .preflight .pf-row .warn { color: #ff6b6b; }

  /* ── Step cards (Setup how-it-works) ── */
  .step-cards { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 16px; margin-bottom: 28px; }
  .step-card { background: #131315; border: 1px solid #1f1f22; border-radius: 8px; padding: 20px 20px 16px; }
  .step-card .sn { font-family: 'JetBrains Mono', monospace; font-size: 10px; color: #D4FF3A; letter-spacing: 0.2em; margin-bottom: 10px; }
  .step-card h5 { font-family: 'Fraunces', serif; font-size: 18px; font-weight: 400; color: #F5F4EE; margin: 0 0 8px 0; }
  .step-card p  { font-family: 'JetBrains Mono', monospace; font-size: 11px; color: #8B8B85; line-height: 1.7; margin: 0; }
  @media (max-width: 768px) {
    .step-cards { grid-template-columns: 1fr; }
    .funnel-row { flex-direction: column; gap: 12px; }
    .funnel-arrow { transform: rotate(90deg); }
  }

  /* ── Status bar (Applied / All Apps) ── */
  .status-bar { display: flex; gap: 8px; flex-wrap: wrap; margin-bottom: 20px; }
  .sb-chip {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    padding: 5px 14px;
    border-radius: 3px;
    letter-spacing: 0.08em;
  }
</style>
""", unsafe_allow_html=True)


# ── Auth wall — must be authenticated to use the app ──────────────
render_auth_wall()
restore_user_session()
_USER_ID = get_user_id()  # set once per Streamlit run, passed to all tracker calls

# ── Tag color patch — JS MutationObserver via same-origin iframe.
#    CSS !important loses to Emotion (BasWeb CSS-in-JS) because Emotion
#    re-injects <style> into <head> after every render. JS setProperty with
#    'important' wins because it writes directly to the element's inline style,
#    which is the highest specificity layer. The observer re-fires on any
#    DOM/attribute mutation so newly-rendered tags are caught immediately.
import streamlit.components.v1 as _components
_components.html("""
<script>
(function(){
  function fix(){
    var tags=window.parent.document.querySelectorAll('span[data-baseweb="tag"]');
    tags.forEach(function(t){
      t.style.setProperty('background','#2a2a2e','important');
      t.style.setProperty('background-color','#2a2a2e','important');
      t.style.setProperty('color','#c8c8c0','important');
      t.style.setProperty('border','1px solid #3e3e44','important');
      t.style.setProperty('border-radius','4px','important');
      t.querySelectorAll('span').forEach(function(s){
        s.style.setProperty('color','#c8c8c0','important');
        s.style.setProperty('background','transparent','important');
        s.style.setProperty('background-color','transparent','important');
      });
      t.querySelectorAll('svg,button,[role="presentation"]').forEach(function(s){
        s.style.setProperty('color','#666','important');
        s.style.setProperty('fill','#666','important');
      });
    });
  }
  fix();
  new MutationObserver(fix).observe(
    window.parent.document.body,
    {childList:true,subtree:true,attributes:true,attributeFilter:['style','class']}
  );
})();
</script>
""", height=0)


# ── Session persistence: restore from ?uid= query param ───────────
def _try_restore_session():
    """On first load, if ?uid= is in the URL, pull resume from Supabase."""
    if st.session_state.get("_session_restored"):
        return  # already ran this run
    st.session_state["_session_restored"] = True

    # Restore the vertical + teaching schedule straight from the URL — this
    # works with zero database columns (survives refresh/bookmark like ?uid=).
    _q_vertical = st.query_params.get("vertical")
    if _q_vertical and not st.session_state.get("vertical"):
        st.session_state["vertical"] = _q_vertical
    _q_sched = st.query_params.get("sched")
    if _q_sched and not st.session_state.get("schedule_pref"):
        st.session_state["schedule_pref"] = _q_sched

    uid = st.query_params.get("uid")
    if not uid:
        return
    if st.session_state.get("resume_text"):
        return  # already have resume in memory

    try:
        data = load_session(uid)
        if data and data.get("resume_text"):
            st.session_state["resume_text"] = data["resume_text"]
            st.session_state["session_uid"] = uid
            if data.get("target_roles"):
                st.session_state["target_roles"] = data["target_roles"]
            if data.get("preferred_locations"):
                st.session_state["preferred_locations"] = data["preferred_locations"]
            if data.get("gmail_scan_enabled"):
                st.session_state["gmail_scan_enabled"] = True
            if data.get("vertical"):
                st.session_state["vertical"] = data["vertical"]
            if data.get("schedule_pref"):
                st.session_state["schedule_pref"] = data["schedule_pref"]
    except Exception:
        pass  # silently skip — DB down or uid not found

_try_restore_session()

# ── Sidebar: Techturi-branded nav ──────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="tt-logo">
      <span class="bracket">[</span>techturi<span class="dot">.</span><span class="bracket">]</span>
    </div>
    <div class="tt-product">Job Pal</div>
    """, unsafe_allow_html=True)

    # Show logged-in user + resume status
    user_email = get_user_email() or ""
    st.markdown(
        f'<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#4A4A45;'
        f'letter-spacing:0.1em;padding:0 20px 4px;overflow:hidden;text-overflow:ellipsis;'
        f'white-space:nowrap" title="{user_email}">{user_email}</div>',
        unsafe_allow_html=True,
    )

    if st.session_state.get("resume_text"):
        st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#D4FF3A;letter-spacing:0.15em;padding:0 20px 12px">✓ RESUME LOADED</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#ff6b6b;letter-spacing:0.15em;padding:0 20px 12px">⚠ NO RESUME — START HERE</div>', unsafe_allow_html=True)

    _nav_map = {
        "Setup":              "Setup",
        "Run Pipeline":       "Run Pipeline",
        "Review Queue":       "Review Queue",
        "Applied":            "Applied",
        "Interviews":         "Interviews",
        "Events":             "Events",
        "Dashboard":          "Dashboard",
        "All Applications":   "All Applications",
        "Assistant":          "Assistant",
    }
    _nav_labels   = list(_nav_map.keys())
    _nav_internal = list(_nav_map.values())
    _forced_page = st.session_state.pop("_nav_page", None)
    if _forced_page and _forced_page in _nav_labels:
        st.session_state["_nav_radio"] = _forced_page
    _sel = st.radio("Navigate", _nav_labels, key="_nav_radio", label_visibility="collapsed")
    page = _nav_map[_sel]

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("Sign Out", use_container_width=True, key="sidebar_signout"):
        sign_out()
        st.rerun()
    st.markdown('<div class="tt-footer" style="margin-top:8px">techturi.org · Tega Eshareturi</div>', unsafe_allow_html=True)


# ── Helpers ────────────────────────────────────────────────────────
def score_badge(score):
    score = score or 0
    cls = "score-high" if score >= 8 else "score-mid" if score >= 6 else "score-low"
    return f'<span class="{cls}">{score}/10</span>'

def status_tag(status):
    cls = f"tag tag-{status}" if status in ("new","applied","interview","rejected","skipped","application_closed","no_response") else "tag"
    label = {"application_closed": "closed", "no_response": "no response"}.get(status, status)
    return f'<span class="{cls}">{label}</span>'

def safe_get_apps():
    """Fetch all applications for the current user."""
    try:
        return get_all_applications(user_id=_USER_ID)
    except Exception as e:
        st.error(f"Cannot connect to database. Check Supabase secrets in Streamlit Cloud settings. Error: `{e}`")
        st.stop()

@st.cache_data(ttl=90, show_spinner=False)
def _cached_review_queue(min_score: int, user_id: str | None) -> list:
    return get_review_queue(min_score=min_score, user_id=user_id)


@st.cache_data(ttl=90, show_spinner=False)
def _cached_personalization_ctx(user_id: str | None) -> dict:
    from tracker import get_personalization_context
    return get_personalization_context(user_id=user_id)


def safe_get_queue(min_score=REVIEW_MIN_SCORE):
    try:
        return _cached_review_queue(min_score, _USER_ID)
    except Exception as e:
        st.error(f"Cannot connect to database. Check Supabase secrets. Error: `{e}`")
        st.stop()

def _days_in_queue(job: dict) -> int | None:
    raw = job.get("created_at")
    if not raw:
        return None
    try:
        created = datetime.fromisoformat(raw)
        if created.tzinfo is None:
            created = created.replace(tzinfo=timezone.utc)
        return (datetime.now(timezone.utc) - created).days
    except Exception:
        return None


def job_card(job, key_prefix, next_statuses, expanded=False):
    """Render a full job card with details, cover letter, and status controls."""
    score = job.get("score") or 0
    icon = "🟢" if score >= 8 else "🟡" if score >= 6 else "🔴"
    days = _days_in_queue(job)
    days_label = f"  ·  {days}d" if days is not None else ""
    with st.expander(
        f"{icon}  {score}/10  —  {job['title'].replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')} @ {job.get('company', '?')}{days_label}",
        expanded=expanded,
    ):
        hint = job.get("_personal_hint")
        bonus = job.get("_personal_bonus") or 0
        if hint:
            # Translate internal reasons into plain English
            friendly = (hint
                .replace("company you engaged positively before", "✅ Company you liked before")
                .replace("company you skipped/rejected before", "⚠️ Company you passed on before")
                .replace("source with past positive outcomes", "📌 Strong source for you")
                .replace("title overlap with roles you liked", "👍 Matches roles you've gone for")
            )
            # Replace "title contains patterns you skip (x, y)" with friendlier form
            friendly = re.sub(
                r"title contains patterns you skip \(([^)]+)\)",
                lambda m: f"⏭️ You tend to skip {m.group(1)} roles",
                friendly
            )
            st.caption(friendly)

        if days is not None:
            color = "#8B8B85" if days <= 7 else "#f5c518" if days <= 14 else "#ff6b6b"
            st.markdown(
                f'<span style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:{color};'
                f'letter-spacing:0.1em">{days}d in queue</span>',
                unsafe_allow_html=True,
            )

        # ⚠️ Type-2 closed heuristic: LinkedIn jobs >14d old with 200+ applicants
        # (guest API cannot detect JS-only closures; this is the best available signal)
        if (days is not None and days > 14
                and "linkedin.com" in (job.get("url") or "").lower()
                and any(p in (job.get("description") or "") for p in
                        ["Over 200 applicants", "200+ applicants", "over 200 applicants"])):
            st.warning("⚠️ Likely filled — LinkedIn shows 200+ applicants and this posting is 2+ weeks old. Verify before applying.")

        col1, col2 = st.columns([2, 1])

        with col1:
            st.markdown(f"**Location:** {job.get('location', 'Unknown')}")
            st.markdown(f"**Score reason:** {job.get('score_reason', '')}")
            # Use stored salary_range first; only re-parse description as fallback
            from agent import _extract_salary_from_text
            salary_range = job.get("salary_range") or _extract_salary_from_text(job.get("description", ""))
            salary_match = job.get("salary_match", "Unknown")
            _is_academic = st.session_state.get("vertical") == "academic"
            if _is_academic:
                # Faculty roles: rank + appointment type matter, not salary bands.
                st.markdown(f"**Rank:** {job.get('seniority', 'Unknown') or 'Unknown'}")
                if job.get("appointment_type"):
                    st.markdown(f"**Appointment:** {job.get('appointment_type')}")
            else:
                if salary_range:
                    match_label = {"Yes": "✅", "No": "❌", "Unknown": "—"}.get(salary_match, "—")
                    st.markdown(f"**Salary:** {salary_range} {match_label}")
                else:
                    st.markdown(f"**Salary:** Not listed — match: {salary_match}")
                st.markdown(f"**Seniority:** {job.get('seniority', 'Unknown')}")
            st.markdown(f"**Source:** {job.get('source', '')}")
            if job.get("url"):
                st.markdown(f"[Open job posting ↗]({job['url']})")

        # Status → outcome mapping: moving to a status auto-logs the matching event
        STATUS_OUTCOME_MAP = {
            "applied":              "applied",
            "interview":            "interview_scheduled",
            "rejected":             "rejected",
            "skipped":              None,
            "application_closed":   None,
            "manual_review":        None,
        }

        with col2:
            new_status = st.selectbox(
                "Move to",
                ["— no change —"] + next_statuses,
                key=f"{key_prefix}_sel_{job['id']}",
            )
            if new_status != "— no change —":
                if st.button("Save", key=f"{key_prefix}_save_{job['id']}", type="secondary"):
                    update_status(job["id"], new_status, user_id=_USER_ID)
                    auto_event = STATUS_OUTCOME_MAP.get(new_status)
                    if auto_event:
                        log_event(job["id"], auto_event, f"status moved to {new_status}", user_id=_USER_ID)
                    else:
                        log_event(job["id"], "status_change", f"{job.get('status', 'unknown')} -> {new_status}", user_id=_USER_ID)
                    # Pop the key so the selectbox resets to default on next render
                    st.session_state.pop(f"{key_prefix}_sel_{job['id']}", None)
                    _cached_review_queue.clear()
                    _cached_personalization_ctx.clear()
                    st.success(f"→ {new_status}")
                    st.rerun()

        # Only show manual signals that have no corresponding status
        st.markdown('<div class="section-label" style="margin-top:16px">Additional Signal</div>', unsafe_allow_html=True)
        signal = st.selectbox(
            "Log signal",
            ["— none —", "recruiter_response", "no_response_14d"],
            key=f"{key_prefix}_signal_{job['id']}",
            help="For events that don't change your status — recruiter reached out, or you've heard nothing after 2 weeks.",
        )
        if signal != "— none —":
            note = st.text_input(
                "Optional note",
                key=f"{key_prefix}_signal_note_{job['id']}",
                placeholder="e.g. recruiter replied in 2 days",
            )
            if st.button("Log Signal", key=f"{key_prefix}_signal_save_{job['id']}"):
                log_event(job["id"], signal, note.strip(), user_id=_USER_ID)
                st.session_state.pop(f"{key_prefix}_signal_{job['id']}", None)
                st.success(f"Logged: {signal}")
                st.rerun()

        # ── Company Research ──────────────────────────────────────
        _company_name = job.get("company", "").strip()
        if _company_name:
            st.markdown('<div class="section-label" style="margin-top:16px">Company Research</div>', unsafe_allow_html=True)
            _show_cr = st.toggle(f"🏢 {_company_name} — history, funding & about", key=f"cr_toggle_{job['id']}", value=False)
            if _show_cr:
                from researcher import research_company, load_cached
                _cr_key = f"_cr_{job['id']}"
                if _cr_key not in st.session_state:
                    _cached = load_cached(_company_name)
                    if _cached:
                        st.session_state[_cr_key] = _cached
                    else:
                        with st.spinner(f"Researching {_company_name}…"):
                            st.session_state[_cr_key] = research_company(
                                _company_name, job_title=job.get("title", "")
                            )

                _cr = st.session_state[_cr_key]
                if _cr.get("summary"):
                    st.markdown(f"**Summary**\n\n{_cr['summary']}")
                if _cr.get("website"):
                    st.markdown(f"**Website:** [{_cr['website']}]({_cr['website']})")
                if _cr.get("funding"):
                    st.markdown(f"**Funding / Valuation**\n\n{_cr['funding']}")
                if _cr.get("about_text"):
                    st.markdown("**About (from their site)**")
                    st.markdown(
                        f'<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;'
                        f'color:#8B8B85;line-height:1.6;padding:10px;background:#131315;'
                        f'border-radius:6px;border:1px solid #1f1f22">{_cr["about_text"][:1200]}</div>',
                        unsafe_allow_html=True,
                    )
                if not any(_cr.get(k) for k in ("summary", "website", "funding", "about_text")):
                    st.caption("No research data found for this company.")
                if st.button("🔄 Refresh research", key=f"cr_refresh_{job['id']}", type="secondary"):
                    from researcher import research_company as _rc_force
                    with st.spinner("Re-fetching…"):
                        st.session_state[_cr_key] = _rc_force(
                            _company_name, job_title=job.get("title", ""), force_refresh=True
                        )
                    st.rerun()

        if job.get("cover_letter"):
            cl_col, dl_col = st.columns([5, 1])
            with cl_col:
                st.markdown("**Cover Letter**")
            with dl_col:
                import io
                from docx import Document
                doc = Document()
                doc.add_heading(f"{job.get('title', '')} — {job.get('company', '')}", level=1)
                for para in job["cover_letter"].split("\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                _resume = st.session_state.get("resume_text") or ""
                _candidate = next((l.strip() for l in _resume.splitlines() if l.strip()), "CoverLetter")
                _candidate = re.sub(r"[^a-zA-Z0-9 ]+", "", _candidate).strip().replace(" ", "_")[:30]
                _co = re.sub(r"[^a-zA-Z0-9]+", "_", job.get("company", "Company")).strip("_")
                st.download_button(
                    "⬇ .docx",
                    data=buf,
                    file_name=f"{_candidate}_{_co}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{job['id']}",
                )
            st.markdown(
                f'<div class="cover-letter">{job["cover_letter"]}</div>',
                unsafe_allow_html=True,
            )
            if st.toggle("📋 Copy-ready (plain text)", key=f"cl_copy_{job['id']}", value=False):
                _cl_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', job["cover_letter"])
                _cl_clean = re.sub(r'\*([^*]+)\*', r'\1', _cl_clean)
                _cl_clean = re.sub(r'\n{3,}', '\n\n', _cl_clean).strip()
                st.code(_cl_clean, language=None)

        # ── Questionnaire Helper (conversational) ─────────────────
        st.markdown('<div class="section-label" style="margin-top:16px">Questionnaire Helper</div>', unsafe_allow_html=True)
        _show_q = st.toggle("📋 Draft answers to application questions", key=f"q_toggle_{job['id']}", value=bool(st.session_state.get(f"q_chat_{job['id']}")))
        if _show_q:
            q_chat_key = f"q_chat_{job['id']}"
            if q_chat_key not in st.session_state:
                st.session_state[q_chat_key] = []

            q_history = st.session_state[q_chat_key]
            resume_text = st.session_state.get("resume_text") or ""

            if not resume_text:
                st.warning("No resume loaded — go to Setup first.")
            else:
                # Pull company research (cached or fetch)
                from researcher import load_cached as _q_load_cr, research_company as _q_rc
                _q_cr_key = f"_cr_{job['id']}"
                if _q_cr_key not in st.session_state:
                    _q_cached = _q_load_cr(job.get("company", ""))
                    st.session_state[_q_cr_key] = _q_cached or {}
                _q_cr = st.session_state.get(_q_cr_key) or {}
                _q_company_block = (
                    f"\nCOMPANY RESEARCH:\n{_q_cr['research_text']}"
                    if _q_cr.get("research_text") else ""
                )

                q_system = [{"type": "text", "text": f"""You are a job application assistant helping the user draft honest, specific answers to application questionnaire questions.

JOB: {job.get('title')} at {job.get('company')}
LOCATION: {job.get('location', 'Not specified')}
WHY IT SCORED WELL: {job.get('score_reason', '')}

JOB DESCRIPTION:
{job.get('description') or ''}
{_q_company_block}

USER RESUME (full):
{resume_text}

{('COVER LETTER (use for tone/framing reference):\\n' + job['cover_letter']) if job.get('cover_letter') else ''}

Instructions:
- Answer each question using specific, honest examples drawn from the resume
- Reference the company's mission/values from the research above when relevant (e.g. "why do you want to work here?")
- Keep each answer 2-4 sentences unless the question clearly needs more
- Format: restate the question briefly, then answer it
- Do not invent experience not in the resume
- You can refine answers on follow-up requests (e.g., "make answer 2 shorter", "add an Oracle example")""", "cache_control": {"type": "ephemeral"}}]

                if not q_history:
                    questions_init = st.text_area(
                        "Paste the application questions",
                        key=f"q_init_{job['id']}",
                        placeholder="1. Why do you want to work at this company?\n2. Describe a time you used data to drive a decision.\n3. ...",
                        height=120,
                    )
                    if st.button("✍️ Draft answers", key=f"q_start_{job['id']}", type="secondary"):
                        if not questions_init.strip():
                            st.warning("Paste some questions first.")
                        else:
                            st.session_state[q_chat_key].append({"role": "user", "content": f"Draft answers to these application questions:\n\n{questions_init.strip()}"})
                            st.rerun()
                else:
                    import anthropic as _anthropic
                    _ac = _anthropic.Anthropic()

                    for msg in q_history:
                        avatar = "👤" if msg["role"] == "user" else "🤖"
                        with st.chat_message(msg["role"], avatar=avatar):
                            st.markdown(msg["content"])

                    if q_history and q_history[-1]["role"] == "user":
                        with st.chat_message("assistant", avatar="🤖"):
                            def _q_gen():
                                with _ac.messages.stream(
                                    model="claude-sonnet-4-6",
                                    max_tokens=2000,
                                    system=q_system,
                                    messages=q_history,
                                    extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
                                ) as stream:
                                    for text in stream.text_stream:
                                        yield text
                            resp = st.write_stream(_q_gen())
                        st.session_state[q_chat_key].append({"role": "assistant", "content": resp})
                        st.rerun()

                    # ── Copy-ready answers ─────────────────────────
                    _last_answer = next(
                        (m["content"] for m in reversed(q_history) if m["role"] == "assistant"),
                        None,
                    )
                    if _last_answer:
                        if st.toggle("📋 Copy-ready answers (plain text)", key=f"q_copy_{job['id']}", value=False):
                            _clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', _last_answer)
                            _clean = re.sub(r'\*([^*]+)\*', r'\1', _clean)
                            _clean = re.sub(r'^#{1,6}\s+', '', _clean, flags=re.MULTILINE)
                            _clean = re.sub(r'\n{3,}', '\n\n', _clean).strip()
                            st.code(_clean, language=None)

                    with st.form(key=f"q_followup_form_{job['id']}", clear_on_submit=True):
                        q_followup = st.text_input(
                            "Follow-up",
                            placeholder="e.g. make answer 2 shorter / add an Oracle example",
                            label_visibility="collapsed",
                        )
                        if st.form_submit_button("Send", type="secondary") and q_followup.strip():
                            st.session_state[q_chat_key].append({"role": "user", "content": q_followup.strip()})
                            st.rerun()

                    if st.button("Start over", key=f"q_clear_{job['id']}", type="secondary"):
                        st.session_state[q_chat_key] = []
                        st.rerun()

def _render_swipe_card_mode(display_queue):
    """Single-card swipe view for the Review Queue. Supports touch swipe + arrow keys."""
    if "rq_passed_ids" not in st.session_state:
        st.session_state["rq_passed_ids"] = set()
    if "rq_actioned_ids" not in st.session_state:
        st.session_state["rq_actioned_ids"] = set()

    passed   = st.session_state["rq_passed_ids"]
    actioned = st.session_state["rq_actioned_ids"]
    effective = [j for j in display_queue if j["id"] not in passed and j["id"] not in actioned]

    total = len(display_queue)
    reviewed = total - len(effective)
    pct = reviewed / total * 100 if total else 100

    st.markdown(
        f'<div class="sc-progress-bar"><div class="sc-progress-fill" style="width:{pct:.0f}%"></div></div>'
        f'<div class="sc-counter">{reviewed} reviewed · {len(effective)} remaining</div>',
        unsafe_allow_html=True,
    )

    if not effective:
        st.markdown("""
        <div class="swipe-card" style="text-align:center;padding:40px 28px">
          <div style="font-size:36px;margin-bottom:12px">✓</div>
          <h4 style="font-family:'Fraunces',serif;font-size:22px;font-weight:300;color:#F5F4EE;margin-bottom:8px">All reviewed</h4>
          <p style="font-family:'JetBrains Mono',monospace;font-size:12px;color:#8B8B85;margin:0">
            You've gone through everything in this batch.
          </p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Reset passes", type="secondary", key="rq_reset_passes"):
            st.session_state["rq_passed_ids"] = set()
            st.rerun()
        return

    job = effective[0]
    score = job.get("score") or 0
    score_cls = "high" if score >= 8 else "mid" if score >= 6 else "low"

    work_type = (job.get("work_type") or "").strip()
    wt_cls = {"Remote": "remote", "Hybrid": "hybrid", "Onsite": "onsite"}.get(work_type, "onsite")
    wt_tag = f'<span class="sc-tag {wt_cls}">{work_type}</span>' if work_type else ""

    days = _days_in_queue(job)
    if days is not None:
        d_cls = "days stale" if days > 14 else "days warn" if days > 7 else "days"
        days_tag = f'<span class="sc-tag {d_cls}">{days}d in queue</span>'
    else:
        days_tag = ""

    from agent import _extract_salary_from_text
    salary = job.get("salary_range") or _extract_salary_from_text(job.get("description", ""))
    salary_tag = f'<span class="sc-tag salary">{salary}</span>' if salary else ""

    hint_html = ""
    if job.get("_personal_hint"):
        hint = (job["_personal_hint"]
            .replace("company you engaged positively before", "✅ Company you liked before")
            .replace("company you skipped/rejected before", "⚠️ Company you passed on before")
            .replace("source with past positive outcomes", "📌 Strong source for you")
            .replace("title overlap with roles you liked", "👍 Matches roles you've gone for"))
        hint = re.sub(
            r"title contains patterns you skip \(([^)]+)\)",
            lambda m: f"⏭️ You tend to skip {m.group(1)} roles",
            hint,
        )
        hint_html = f'<div class="sc-hint">{hint}</div>'

    url = job.get("url", "")
    link_html = f'<a class="sc-link" href="{url}" target="_blank" rel="noopener">View posting ↗</a>' if url else ""

    title_safe = job.get("title", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    company = job.get("company", "?")
    location = job.get("location", "")
    reason = job.get("score_reason", "")

    st.markdown(f"""
    <div class="swipe-card" id="jp-swipe-card">
      <span class="sc-swipe-indicator skip-ind"  id="jp-skip-ind">← SKIP</span>
      <span class="sc-swipe-indicator apply-ind" id="jp-apply-ind">APPLY →</span>
      <div class="sc-score-badge {score_cls}">{score}<span style="font-size:20px;color:#4A4A45">/10</span></div>
      <div class="sc-score-label">AI match score</div>
      <h2 class="sc-title">{title_safe}</h2>
      <div class="sc-company">{company} · {location}</div>
      <div class="sc-tags">{wt_tag}{salary_tag}{days_tag}</div>
      <div class="sc-reason">{reason}</div>
      {hint_html}
      {link_html}
    </div>
    """, unsafe_allow_html=True)

    # Swipe + keyboard JS — idempotent (window flag prevents double-attach across rerenders)
    st.markdown("""
    <script>
    (function() {
      if (window._jpSwipeInit) return;
      window._jpSwipeInit = true;

      function clickBtn(text) {
        var btns = document.querySelectorAll('button');
        for (var i = 0; i < btns.length; i++) {
          if (btns[i].innerText.trim() === text) { btns[i].click(); return; }
        }
      }

      // Arrow keys (skip focus on inputs)
      document.addEventListener('keydown', function(e) {
        var tag = document.activeElement ? document.activeElement.tagName : '';
        if (tag === 'INPUT' || tag === 'TEXTAREA') return;
        if (e.key === 'ArrowLeft')  { e.preventDefault(); clickBtn('← Skip'); }
        if (e.key === 'ArrowRight') { e.preventDefault(); clickBtn('Apply →'); }
        if (e.key === 'ArrowDown')  { e.preventDefault(); clickBtn('Pass'); }
      });

      // Touch swipe with live card drag + indicators
      var tx = 0, ty = 0, dragging = false;
      document.addEventListener('touchstart', function(e) {
        tx = e.touches[0].clientX;
        ty = e.touches[0].clientY;
        dragging = false;
      }, { passive: true });

      document.addEventListener('touchmove', function(e) {
        var dx = e.touches[0].clientX - tx;
        var dy = e.touches[0].clientY - ty;
        if (!dragging && Math.abs(dx) < Math.abs(dy)) return; // vertical scroll
        dragging = true;
        var card = document.getElementById('jp-swipe-card');
        if (!card) return;
        card.style.transform = 'translateX(' + (dx * 0.35) + 'px) rotate(' + (dx * 0.018) + 'deg)';
        var si = document.getElementById('jp-skip-ind');
        var ai = document.getElementById('jp-apply-ind');
        if (si)  si.style.opacity  = dx < -20 ? Math.min(1, (-dx - 20) / 70).toFixed(2) : 0;
        if (ai) ai.style.opacity  = dx > 20  ? Math.min(1, (dx - 20) / 70).toFixed(2)  : 0;
      }, { passive: true });

      document.addEventListener('touchend', function(e) {
        var card = document.getElementById('jp-swipe-card');
        if (card) card.style.transform = '';
        var si = document.getElementById('jp-skip-ind');
        var ai = document.getElementById('jp-apply-ind');
        if (si) si.style.opacity  = 0;
        if (ai) ai.style.opacity = 0;
        if (!dragging) return;
        var dx = e.changedTouches[0].clientX - tx;
        if (Math.abs(dx) > 75) {
          if (dx < 0) clickBtn('← Skip');
          else        clickBtn('Apply →');
        }
      }, { passive: true });
    })();
    </script>
    """, unsafe_allow_html=True)

    skip_col, pass_col, apply_col = st.columns([5, 3, 5])
    with skip_col:
        if st.button("← Skip", use_container_width=True, key=f"cs_{job['id']}"):
            update_status(job["id"], "skipped", user_id=_USER_ID)
            log_event(job["id"], "status_change", "card skip", user_id=_USER_ID)
            actioned.add(job["id"])
            passed.discard(job["id"])
            st.rerun()
    with pass_col:
        if st.button("Pass", use_container_width=True, key=f"cp_{job['id']}"):
            passed.add(job["id"])
            st.rerun()
    with apply_col:
        if st.button("Apply →", use_container_width=True, type="primary", key=f"ca_{job['id']}"):
            update_status(job["id"], "applied", user_id=_USER_ID)
            log_event(job["id"], "applied", "card swipe", user_id=_USER_ID)
            actioned.add(job["id"])
            passed.discard(job["id"])
            st.rerun()

    st.markdown(
        '<div class="sc-gesture-hint">← skip · pass · apply →  ·  arrow keys on desktop</div>',
        unsafe_allow_html=True,
    )


def _categorize(title: str) -> str:
    t = (title or "").lower()
    if any(k in t for k in ["genai", "gen ai", "generative", "llm", "ai engineer", "ml engineer", "machine learning"]):
        return "AI / ML"
    if any(k in t for k in ["data engineer", "etl", "pipeline", "databricks", "spark", "platform engineer"]):
        return "Data Engineering"
    if any(k in t for k in ["analytics engineer", "data analyst", "business intelligence", "bi engineer", "data scientist"]):
        return "Analytics"
    if any(k in t for k in ["program manager", "project manager", "tpm", "technical program", "scrum"]):
        return "Program Management"
    if any(k in t for k in ["software engineer", "backend", "frontend", "full stack", "fullstack", "swe"]):
        return "Software Engineering"
    if any(k in t for k in ["data architect", "cloud architect", "solutions architect"]):
        return "Architecture"
    return "Other"


def page_header(eyebrow, title_html):
    st.markdown(f"""
    <div class="page-header">
      <div>
        <div class="eyebrow">{eyebrow}</div>
        <h1>{title_html}</h1>
      </div>
    </div>
    """, unsafe_allow_html=True)

def _build_assistant_system_prompt(page_context: str = "") -> str:
    """Build the shared assistant system prompt with live pipeline context."""
    apps = []
    try:
        apps = get_all_applications(user_id=_USER_ID) or []
    except Exception:
        pass
    resume_text = st.session_state.get("resume_text") or ""
    target_roles = st.session_state.get("target_roles") or []
    min_salary   = st.session_state.get("min_salary") or 0

    new_jobs   = [a for a in apps if a.get("status") == "new"]
    applied    = [a for a in apps if a.get("status") == "applied"]
    interviews = [a for a in apps if a.get("status") == "interview"]

    strong   = sorted([j for j in new_jobs if (j.get("score") or 0) >= 8], key=lambda x: x.get("score") or 0, reverse=True)
    solid    = sorted([j for j in new_jobs if (j.get("score") or 0) == 7], key=lambda x: x.get("score") or 0, reverse=True)
    consider = sorted([j for j in new_jobs if (j.get("score") or 0) == 6], key=lambda x: x.get("score") or 0, reverse=True)

    def _job_line(a):
        sal = a.get("salary_range") or "no salary"
        return f"- {a.get('title')} @ {a.get('company')} | score {a.get('score')}/10 | {sal}"

    strong_lines  = "\n".join(_job_line(j) for j in strong[:8])   or "None"
    solid_lines   = "\n".join(_job_line(j) for j in solid[:8])    or "None"
    consider_lines= "\n".join(_job_line(j) for j in consider[:8]) or "None"
    applied_lines = "\n".join(
        f"- {a.get('title')} @ {a.get('company')}" for a in applied[:10]
    ) or "None yet"
    interview_lines = "\n".join(
        f"- {a.get('title')} @ {a.get('company')}" for a in interviews
    ) or "None"

    context_block = f"\nCURRENT PAGE: {page_context}\n" if page_context else ""

    return f"""You are Job Pal Assistant, an AI job search co-pilot built into the Job Pal platform.
You have full context on this user's resume, job queue, and application history.
{context_block}
USER PROFILE
Target roles: {', '.join(target_roles) or 'not set'}
Minimum salary: {'${:,}'.format(min_salary) if min_salary else 'not set'}

FULL RESUME
{resume_text}

JOB QUEUE — UNREVIEWED JOBS (status: new)

STRONG LEADS — score 8-10 ({len(strong)} total, showing top 8)
{strong_lines}

SOLID LEADS — score 7 ({len(solid)} total, showing top 8)
{solid_lines}

WORTH CONSIDERING — score 6 ({len(consider)} total, showing top 8)
{consider_lines}

APPLIED ({len(applied)} total, showing 10 most recent)
{applied_lines}

INTERVIEWS ({len(interviews)} active)
{interview_lines}

TOTAL TRACKED: {len(apps)} jobs | Unreviewed: {len(new_jobs)} | Applied: {len(applied)}

SCORE GUIDE: 8-10 = strong match (apply now), 7 = solid (worth applying), 6 = partial fit (apply if targeting this role), 5 and below = long shot.

You help the user with:
- Deciding which queue jobs to apply to first (recommend 8s first, then 7s, then 6s if quota allows)
- Understanding why a job scored the way it did
- Writing or refining cover letters, thank-you emails, follow-ups
- Interview prep for companies in their pipeline
- Resume gap analysis against specific roles
- Salary negotiation guidance
- General job search strategy

Be direct, specific, and use the data above. When recommending jobs from the queue, always mention the score and company name.
Keep answers concise unless asked for detail. Never invent job details — only reference what's shown above.
Do not give legal or financial advice."""


def _build_dynamic_suggestions(apps: list[dict]) -> list[str]:
    """Generate contextual starter suggestions based on the live pipeline."""
    interviews  = [a for a in apps if a.get("status") == "interview"]
    applied     = [a for a in apps if a.get("status") == "applied"]
    new_jobs    = [a for a in apps if a.get("status") == "new"]
    strong      = [j for j in new_jobs if (j.get("score") or 0) >= 8]

    suggestions = []

    # Interview prep — most urgent, show first
    for iv in interviews[:2]:
        suggestions.append(
            f"Help me prep for my {iv.get('company')} interview — what should I focus on?"
        )

    # Strong leads if any
    if strong:
        top = strong[0]
        suggestions.append(
            f"I have a {top.get('score')}/10 match at {top.get('company')} — help me tailor my application."
        )

    # Applied follow-up
    if applied:
        suggestions.append("Draft a follow-up email for a job I applied to last week.")

    # Always-useful fallbacks
    fallbacks = [
        "Which jobs in my queue should I apply to first?",
        "What skills am I missing for the roles I'm targeting?",
        "Which of my 6-scoring jobs are still worth applying to?",
        "How should I position my Oracle experience for data engineering roles?",
    ]
    for f in fallbacks:
        if len(suggestions) >= 4:
            break
        if f not in suggestions:
            suggestions.append(f)

    return suggestions[:4]


def metric(col, label, value, sub="", accent=False):
    col.markdown(f"""
    <div class="metric-card{"  accent" if accent else ""}">
      <div class="label">{label}</div>
      <div class="value">{value}</div>
      {"<div class='sub'>" + sub + "</div>" if sub else ""}
    </div>""", unsafe_allow_html=True)


def _boolish(v) -> bool:
    if isinstance(v, bool):
        return v
    if v is None:
        return False
    return str(v).lower() in ("true", "1", "t", "yes")


def enrich_experiment_runs_df(df: pd.DataFrame) -> pd.DataFrame:
    """Add qualified_pct and human-readable cost_mode for experiment comparison."""
    if df.empty:
        return df
    out = df.copy()

    def qualified_pct_row(row):
        new = int(row.get("jobs_new") or 0)
        q = int(row.get("jobs_qualified") or 0)
        return round(100.0 * q / new, 1) if new else None

    def cost_mode_row(row):
        mode = str(row.get("scoring_mode") or "").lower().strip()
        letters = _boolish(row.get("cover_letters_enabled"))
        if mode == "cheap":
            return "No LLM"
        if mode == "hybrid":
            return "Hybrid + letters" if letters else "Hybrid (score only)"
        return "Claude + letters" if letters else "Claude (score only)"

    out["qualified_pct"] = out.apply(qualified_pct_row, axis=1)
    out["cost_mode"] = out.apply(cost_mode_row, axis=1)
    return out


# ═══════════════════════════════════════════════════════════════════
# DASHBOARD
# ═══════════════════════════════════════════════════════════════════
BETA_JOB_LIMIT = 50  # max jobs scored per pipeline run for beta users

# ═══════════════════════════════════════════════════════════════════
# SETUP — resume onboarding (required before pipeline runs)
# ═══════════════════════════════════════════════════════════════════
if page == "Setup":
    page_header("Setup", "Start <em>here.</em>")

    st.markdown("""
    <div class="step-cards">
      <div class="step-card">
        <div class="sn">01 — Resume</div>
        <h5>Tell it who you are</h5>
        <p>Upload a PDF/DOCX or paste your resume. Job Pal reads every line to score job fit 1–10 against your actual background.</p>
      </div>
      <div class="step-card">
        <div class="sn">02 — Pipeline</div>
        <h5>Let AI do the hunting</h5>
        <p>Scrapes LinkedIn, Remotive, RemoteOK & more in parallel. Claude scores every listing and writes a tailored cover letter for strong matches.</p>
      </div>
      <div class="step-card">
        <div class="sn">03 — Review</div>
        <h5>Apply to the best</h5>
        <p>Your top matches land in Review Queue. Move them through Applied → Interviews. Track response rates and outcomes on the Dashboard.</p>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Resume status + start-fresh action ───────────────────────
    if st.session_state.get("resume_text"):
        status_col, clear_col = st.columns([3, 1])
        with status_col:
            st.success("✓ Resume loaded — update below and hit Save to switch resumes.")
        with clear_col:
            if st.button("Start Fresh", use_container_width=True, key="setup_start_fresh"):
                if st.session_state.get("_confirm_setup_clear"):
                    from tracker import clear_all_data
                    clear_all_data(user_id=_USER_ID)
                    if _USER_ID:
                        clear_session(_USER_ID)
                    for _k in ("resume_text", "target_roles", "min_salary", "session_uid",
                               "preferred_locations", "_suggested_roles", "_roles_text_area",
                               "_confirm_setup_clear", "_session_restored"):
                        st.session_state.pop(_k, None)
                    st.session_state["_user_session_restored_for"] = _USER_ID
                    st.query_params.clear()
                    st.rerun()
                else:
                    st.session_state["_confirm_setup_clear"] = True
        if st.session_state.get("_confirm_setup_clear"):
            st.warning("This clears all jobs, pipeline runs, and your saved resume. Click **Start Fresh** again to confirm.")
    else:
        st.info("Upload or paste your resume below, then hit **Save**.")

    # ── Vertical: what kind of roles are you looking for? ─────────
    st.markdown('<div class="section-label" style="margin-top:20px">What are you looking for?</div>', unsafe_allow_html=True)
    _vert_labels = {
        "tech":     "💻  Tech & Business roles",
        "academic": "🎓  Faculty & Adjunct (Higher Ed)",
    }
    _vert_keys = list(_vert_labels.keys())
    _cur_vert = st.session_state.get("vertical", "tech")
    _sel_vert_label = st.radio(
        "Vertical",
        [_vert_labels[k] for k in _vert_keys],
        index=_vert_keys.index(_cur_vert) if _cur_vert in _vert_keys else 0,
        horizontal=True,
        label_visibility="collapsed",
    )
    _selected_vertical = _vert_keys[[_vert_labels[k] for k in _vert_keys].index(_sel_vert_label)]
    st.session_state["vertical"] = _selected_vertical

    # ── Teaching availability (academic only) ────────────────────
    _schedule_pref = st.session_state.get("schedule_pref", "")
    if _selected_vertical == "academic":
        st.caption("Faculty mode: Job Pal scores postings on degree-in-field, discipline and teaching fit — and matches your class-schedule availability for adjunct sections.")
        _sched_options = ["Evening / night classes", "Early-morning classes", "Weekend classes", "Online / asynchronous"]
        _saved_sched = [s.strip() for s in (_schedule_pref.split(",") if _schedule_pref else []) if s.strip()]
        _sched_selected = st.multiselect(
            "Preferred teaching schedule",
            _sched_options,
            default=[s for s in _saved_sched if s in _sched_options],
            help="Adjunct sections are often evenings/weekends — pick what fits around your schedule.",
        )
        _schedule_pref = ", ".join(_sched_selected)
        st.session_state["schedule_pref"] = _schedule_pref

    # ── Database connection diagnostic ───────────────────────────
    # Answers "which Supabase am I connected to?" and, when the academic
    # columns are missing, shows the exact SQL to run in THAT project.
    with st.expander("🔌 Database connection", expanded=False):
        try:
            from tracker import session_columns_status
            _st = session_columns_status()
            st.caption(f"Connected Supabase project: **`{_st['project_ref']}`**  "
                       f"(run any migration in *this* project's SQL editor).")
            if not _st["has_user_sessions"]:
                st.warning("`user_sessions` table not found in this project — you may be "
                           "pointed at the wrong Supabase project.")
            elif _st["has_vertical"] and _st["has_schedule_pref"]:
                st.success("Faculty-mode columns present — your vertical & schedule persist across refreshes.")
            else:
                st.info("Faculty mode works now, but to make it persist across refreshes, "
                        "run this once in the SQL editor of the project above:")
                st.code(
                    "alter table user_sessions add column if not exists vertical text default 'tech';\n"
                    "alter table user_sessions add column if not exists schedule_pref text;",
                    language="sql",
                )
                st.caption("(Can't be automated — the anon key can't run ALTER TABLE.)")
        except Exception as _e:
            st.caption(f"Could not check database connection: {_e}")

    # ── Resume input: upload or paste ────────────────────────────
    st.markdown('<div class="section-label">Your Resume</div>', unsafe_allow_html=True)
    tab_upload, tab_paste = st.tabs(["Upload File", "Paste Text"])

    uploaded_text = None   # from file upload
    pasted_text   = None   # from paste tab

    with tab_upload:
        uploaded = st.file_uploader(
            "Upload your resume",
            type=["pdf", "docx", "txt"],
            label_visibility="collapsed",
            help="PDF, DOCX, or TXT — text is extracted automatically",
        )
        if uploaded:
            try:
                if uploaded.type == "application/pdf":
                    import fitz  # PyMuPDF
                    doc = fitz.open(stream=uploaded.read(), filetype="pdf")
                    uploaded_text = "\n".join(p.get_text() for p in doc)
                elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    from docx import Document
                    doc = Document(uploaded)
                    uploaded_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                else:
                    uploaded_text = uploaded.read().decode("utf-8", errors="ignore")

                st.success(f"✓ {uploaded.name} extracted — {len(uploaded_text)} characters")
                st.markdown(
                    f'<div class="cover-letter" style="max-height:200px;overflow-y:auto">{uploaded_text[:800]}{"..." if len(uploaded_text) > 800 else ""}</div>',
                    unsafe_allow_html=True,
                )
            except Exception as e:
                st.error(f"Could not read file: {e}")

    with tab_paste:
        pasted = st.text_area(
            "Paste resume text",
            value=st.session_state.get("resume_text") or "",
            height=280,
            placeholder="Paste your full resume as plain text — work history, skills, education, certifications...",
            label_visibility="collapsed",
        )
        if pasted.strip():
            pasted_text = pasted.strip()

    # Upload takes priority over paste; paste is fallback
    extracted_text = uploaded_text or pasted_text

    # ── Salary preference ────────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:20px">Minimum Salary (optional)</div>', unsafe_allow_html=True)
    salary_options = [
        "No minimum",
        "$80,000+",
        "$100,000+",
        "$120,000+",
        "$140,000+",
        "$160,000+",
        "$180,000+",
        "$200,000+",
        "$220,000+",
    ]
    salary_map = {
        "No minimum": 0,
        "$80,000+": 80_000,
        "$100,000+": 100_000,
        "$120,000+": 120_000,
        "$140,000+": 140_000,
        "$160,000+": 160_000,
        "$180,000+": 180_000,
        "$200,000+": 200_000,
        "$220,000+": 220_000,
    }
    current_min = st.session_state.get("min_salary", 0)
    current_label = next((k for k, v in salary_map.items() if v == current_min), "No minimum")
    selected_salary = st.selectbox(
        "Minimum salary",
        salary_options,
        index=salary_options.index(current_label),
        label_visibility="collapsed",
    )

    # ── Target roles ─────────────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:20px">Target Roles — one per line</div>', unsafe_allow_html=True)

    if _selected_vertical == "academic":
        from config import ACADEMIC_ROLES as _DEFAULT_ROLES
    else:
        from config import TARGET_ROLES as _DEFAULT_ROLES
    saved_roles = st.session_state.get("target_roles") or _DEFAULT_ROLES
    default_roles = "\n".join(saved_roles)

    suggest_col, _ = st.columns([1, 3])
    with suggest_col:
        suggest_btn = st.button(
            "Suggest from resume",
            disabled=not bool(extracted_text),
            help="Uses AI to infer target job titles from your resume. Paste or upload your resume first.",
        )

    if suggest_btn and extracted_text:
        with st.spinner("Reading your resume..."):
            try:
                import anthropic as _ant
                from config import ANTHROPIC_API_KEY
                _ac = _ant.Anthropic(api_key=ANTHROPIC_API_KEY)

                # Capture whatever the user currently has in the text area
                current_roles_raw = st.session_state.get("_roles_text_area") or "\n".join(
                    st.session_state.get("target_roles") or []
                )
                current_roles = [r.strip() for r in current_roles_raw.splitlines() if r.strip()]

                if current_roles:
                    existing_block = "ROLES THE USER HAS ALREADY SET:\n" + "\n".join(current_roles) + "\n\n"
                    instruction = (
                        "The user has already manually selected the roles above. "
                        "Treat those as high-priority — keep every one that is a strong fit for this resume. "
                        "Add any additional titles from the resume that are missing. "
                        "Only remove a manually-set role if it is a genuinely poor fit (wrong field entirely). "
                        "Output the final merged list as job titles only, one per line, no bullets, no explanations. "
                        "Include seniority level where appropriate. 10–15 titles total.\n\n"
                    )
                else:
                    existing_block = ""
                    instruction = (
                        "List 10–15 specific job titles this person should target in their job search. "
                        "Output ONLY the job titles, one per line, no bullets, no explanations, no numbering. "
                        "Include seniority level where appropriate.\n\n"
                    )

                _msg = _ac.messages.create(
                    model="claude-haiku-4-5",
                    max_tokens=400,
                    messages=[{
                        "role": "user",
                        "content": (
                            f"{existing_block}{instruction}"
                            f"RESUME:\n{extracted_text[:3000]}"
                        ),
                    }],
                )
                st.session_state["_suggested_roles"] = _msg.content[0].text.strip()
                st.session_state["_roles_text_area"] = st.session_state["_suggested_roles"]
            except Exception as e:
                st.warning(f"Could not suggest roles: {e}")

    # Read suggested roles without popping — pop caused the text area to revert on next interaction
    role_value = st.session_state.get("_roles_text_area") or default_roles

    roles_input = st.text_area(
        "Target roles",
        value=role_value,
        height=180,
        placeholder="Paste your resume above first, then click Suggest — or type roles manually.\nOne role per line.",
        label_visibility="collapsed",
    )

    # Show extracted skills so user can verify the system understands their background
    if extracted_text and len(extracted_text.strip()) > 100:
        try:
            from agent import _extract_skills
            detected = sorted(_extract_skills(extracted_text))
            if detected:
                with st.expander(f"Skills detected from your resume ({len(detected)} found)", expanded=False):
                    st.caption("These are the skills Job Pal will use to score job fit. If key skills are missing, make sure they appear in your resume text.")
                    st.markdown(" ".join(
                        f'<span style="background:#1C1C18;border:1px solid #2A2A25;border-radius:4px;'
                        f'padding:2px 8px;font-family:\'JetBrains Mono\',monospace;font-size:11px;'
                        f'color:#D4FF3A;margin:2px;display:inline-block">{s}</span>'
                        for s in detected
                    ), unsafe_allow_html=True)
        except Exception:
            pass

    # ── Preferred Locations ───────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:20px">Preferred Locations</div>', unsafe_allow_html=True)
    _loc_options = [
        # ── Special ──────────────────────────────────────────────
        "Remote",
        # ── United States ─────────────────────────────────────────
        "Atlanta, GA", "Austin, TX", "Baltimore, MD", "Boston, MA",
        "Charlotte, NC", "Chicago, IL", "Cincinnati, OH", "Cleveland, OH",
        "Columbus, OH", "Dallas, TX", "Denver, CO", "Detroit, MI",
        "El Paso, TX", "Fort Worth, TX", "Houston, TX", "Indianapolis, IN",
        "Jacksonville, FL", "Kansas City, MO", "Las Vegas, NV", "Los Angeles, CA",
        "Louisville, KY", "Memphis, TN", "Miami, FL", "Milwaukee, WI",
        "Minneapolis, MN", "Nashville, TN", "New Orleans, LA", "New York, NY",
        "Oklahoma City, OK", "Orlando, FL", "Philadelphia, PA", "Phoenix, AZ",
        "Pittsburgh, PA", "Portland, OR", "Raleigh, NC", "Richmond, VA",
        "Sacramento, CA", "Salt Lake City, UT", "San Antonio, TX", "San Diego, CA",
        "San Francisco, CA", "San Jose, CA", "Seattle, WA", "St. Louis, MO",
        "Tampa, FL", "Tucson, AZ", "Virginia Beach, VA", "Washington, DC",
        # ── Canada ────────────────────────────────────────────────
        "Toronto, Canada", "Vancouver, Canada", "Montreal, Canada",
        "Calgary, Canada", "Ottawa, Canada", "Edmonton, Canada",
        # ── United Kingdom ────────────────────────────────────────
        "London, UK", "Manchester, UK", "Birmingham, UK", "Edinburgh, UK",
        # ── Europe ────────────────────────────────────────────────
        "Amsterdam, Netherlands", "Barcelona, Spain", "Berlin, Germany",
        "Dublin, Ireland", "Frankfurt, Germany", "Lisbon, Portugal",
        "Madrid, Spain", "Munich, Germany", "Paris, France", "Zurich, Switzerland",
        # ── Asia-Pacific ──────────────────────────────────────────
        "Bangalore, India", "Dubai, UAE", "Hong Kong", "Hyderabad, India",
        "Melbourne, Australia", "Mumbai, India", "Singapore",
        "Sydney, Australia", "Tokyo, Japan",
        # ── Latin America ─────────────────────────────────────────
        "Bogotá, Colombia", "Buenos Aires, Argentina", "Mexico City, Mexico",
        "São Paulo, Brazil",
    ]
    _saved_locs = st.session_state.get("preferred_locations") or ["Remote"]
    # Saved locations may include cities not in the preset list — merge them in so they stay selected
    _extra_saved = [l for l in _saved_locs if l not in _loc_options]
    _full_options = _loc_options + _extra_saved
    selected_locations = st.multiselect(
        "Preferred locations",
        options=_full_options,
        default=[l for l in _saved_locs if l in _full_options],
        label_visibility="collapsed",
        placeholder="Type to search cities or countries…",
        help="Start typing any city or country — the list filters as you type. Select Remote to include remote-only boards.",
    )

    if st.button("Save & Go to Pipeline →", type="primary", use_container_width=True):
        if not extracted_text or len(extracted_text.strip()) < 100:
            st.error("Resume looks too short or empty — upload a file or paste your resume text.")
        else:
            clean_text = extracted_text.strip()
            roles = [r.strip() for r in roles_input.splitlines() if r.strip()]
            locs  = selected_locations or ["Remote"]
            st.session_state["resume_text"] = clean_text
            st.session_state["target_roles"] = roles
            st.session_state["min_salary"] = salary_map[selected_salary]
            st.session_state["preferred_locations"] = locs
            st.session_state["vertical"] = _selected_vertical
            st.session_state["schedule_pref"] = _schedule_pref

            # Persist to Supabase — use auth user_id as key so resume is tied to account
            uid = _USER_ID or st.session_state.get("session_uid") or new_uid()
            # Persist vertical + schedule in the URL so faculty mode survives a
            # refresh/bookmark even when the DB columns aren't present.
            st.query_params["vertical"] = _selected_vertical
            if _schedule_pref:
                st.query_params["sched"] = _schedule_pref
            elif "sched" in st.query_params:
                del st.query_params["sched"]
            try:
                save_session(uid, clean_text, roles, preferred_locations=locs,
                             vertical=_selected_vertical, schedule_pref=_schedule_pref)
                st.session_state["session_uid"] = uid
                if not _USER_ID:
                    st.query_params["uid"] = uid
                st.success(
                    f"✓ Resume saved. Bookmark this page — your resume will reload automatically. "
                    f"Click **Run Pipeline** to start."
                )
            except Exception as e:
                # DB write failed — session still works in memory this tab
                st.success("✓ Resume saved for this session. Click **Run Pipeline** to start.")
                st.caption(f"(Persistence unavailable: {e})")

    if st.session_state.get("resume_text"):
        st.markdown('<div class="section-label" style="margin-top:28px">Currently Loaded Resume</div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="cover-letter" style="max-height:200px;overflow-y:auto">{st.session_state["resume_text"][:800]}{"..." if len(st.session_state["resume_text"]) > 800 else ""}</div>',
            unsafe_allow_html=True,
        )
        st.caption("To update, upload a new file or paste new text and hit Save again.")

    # ── Gmail Rejection Scanning opt-in ──────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:36px">Gmail Rejection Scanning</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="pipeline-card" style="margin-bottom:16px">
      <p style="margin:0">
        <b>Optional.</b> Connect your job-application inbox so Job Pal can scan for
        rejection emails. When enabled, the Applied page scans automatically the first
        time you visit each session. Your credentials are <b>never stored</b> — enter
        them once per browser session and Job Pal handles the rest.
      </p>
    </div>
    """, unsafe_allow_html=True)

    _gmail_enabled = st.session_state.get("gmail_scan_enabled", False)
    _gmail_toggle = st.toggle(
        "Enable Gmail rejection scanning",
        value=_gmail_enabled,
        key="gmail_opt_in_toggle",
    )
    if _gmail_toggle != _gmail_enabled:
        st.session_state["gmail_scan_enabled"] = _gmail_toggle
        if _USER_ID:
            save_gmail_opt_in(_USER_ID, _gmail_toggle)

    if _gmail_toggle:
        st.markdown("""
        <div style="font-family:'JetBrains Mono',monospace;font-size:11px;color:#8B8B85;margin-bottom:8px">
        You need a <b>Gmail App Password</b> — not your regular password.<br>
        Create one at: <b>myaccount.google.com → Security → 2-Step Verification → App passwords</b>
        </div>
        """, unsafe_allow_html=True)

        _gcol1, _gcol2 = st.columns(2)
        with _gcol1:
            _gmail_email = st.text_input(
                "Gmail address you apply with",
                value=st.session_state.get("gmail_apply_email", "tegapeters11@gmail.com"),
                key="gmail_email_input",
            )
        with _gcol2:
            _gmail_pw = st.text_input(
                "App Password (16-char, no spaces)",
                type="password",
                placeholder="xxxx xxxx xxxx xxxx",
                key="gmail_pw_input",
            )
        if st.button("Save Gmail credentials for this session", key="save_gmail_creds"):
            if _gmail_email and _gmail_pw:
                st.session_state["gmail_apply_email"] = _gmail_email.strip()
                st.session_state["gmail_app_password"] = _gmail_pw.replace(" ", "").strip()
                st.success("✓ Gmail credentials saved. The Applied page will scan automatically when you visit it.")
            else:
                st.warning("Enter both email and app password.")

elif page == "Dashboard":
    page_header("Dashboard", "Your <em>pipeline.</em>")

    apps = safe_get_apps()
    if not apps:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:48px 24px">
          <h4>Nothing tracked yet</h4>
          <p>Run the pipeline to start scoring jobs against your resume.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("▶ Run Pipeline", type="primary"):
            st.session_state["_nav_page"] = "Run Pipeline"
            st.rerun()
        st.stop()

    df = pd.DataFrame(apps)
    status_counts = df["status"].value_counts().to_dict()
    n_queue    = status_counts.get("new", 0)
    n_applied  = status_counts.get("applied", 0)
    n_interview= status_counts.get("interview", 0)
    n_rejected = status_counts.get("rejected", 0)
    avg_score  = df["score"].dropna().mean()
    response_rate = round(n_interview / n_applied * 100) if n_applied else 0

    # ── Metrics row ───────────────────────────────────────────────
    col1, col2, col3, col4, col5 = st.columns(5)
    metric(col1, "Total Tracked",  len(df))
    metric(col2, "In Queue",       n_queue,     "awaiting review")
    metric(col3, "Applied",        n_applied)
    metric(col4, "Interviews",     n_interview, "active", accent=True)
    metric(col5, "Avg AI Score",   f"{avg_score:.1f}" if not df["score"].dropna().empty else "—", "out of 10")

    # ── Funnel ────────────────────────────────────────────────────
    n_qualified = len(df[df["score"].fillna(0) >= REVIEW_MIN_SCORE])
    st.markdown(f"""
    <div class="funnel-row">
      <div class="funnel-step">
        <div class="fn">{len(df)}</div>
        <div class="fl">Scraped</div>
      </div>
      <div class="funnel-arrow">→</div>
      <div class="funnel-step">
        <div class="fn accent">{n_qualified}</div>
        <div class="fl">Scored {REVIEW_MIN_SCORE}+</div>
      </div>
      <div class="funnel-arrow">→</div>
      <div class="funnel-step">
        <div class="fn">{n_applied}</div>
        <div class="fl">Applied</div>
      </div>
      <div class="funnel-arrow">→</div>
      <div class="funnel-step">
        <div class="fn {'accent' if n_interview else ''}">{n_interview}</div>
        <div class="fl">Interviews</div>
      </div>
      <div class="funnel-arrow">→</div>
      <div class="funnel-step">
        <div class="fn {'accent' if response_rate >= 20 else ''}">{response_rate}%</div>
        <div class="fl">Response Rate</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Analytics row ─────────────────────────────────────────────
    st.markdown('<div class="section-label">Analytics</div>', unsafe_allow_html=True)
    col_a, col_b = st.columns(2)
    with col_a:
        st.caption("Score distribution")
        score_df = df["score"].dropna().value_counts().sort_index()
        if not score_df.empty:
            st.bar_chart(score_df, color="#D4FF3A")
    with col_b:
        st.caption("Status breakdown")
        nice_status = {k: v for k, v in status_counts.items() if v > 0}
        if nice_status:
            st.bar_chart(pd.Series(nice_status), color="#D4FF3A")

    event_counts = get_event_counts(user_id=_USER_ID)
    if event_counts:
        st.markdown('<div class="section-label" style="margin-top:8px">Outcome signals</div>', unsafe_allow_html=True)
        st.bar_chart(pd.Series(event_counts), color="#D4FF3A")

    # ── Source health ─────────────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:8px">Source health</div>', unsafe_allow_html=True)
    st.caption("Volume, avg score, share scoring 7+, and how many are waiting in Review Queue — per job board.")
    health_df = pd.DataFrame(get_source_health(apps=apps, review_min_score=REVIEW_MIN_SCORE, user_id=_USER_ID))
    if not health_df.empty:
        st.dataframe(health_df, use_container_width=True, hide_index=True)

    # ── Top matches ───────────────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:8px">Top matches</div>', unsafe_allow_html=True)
    top_cols = ["title", "company", "location", "score", "score_reason", "status"]
    top = df.nlargest(10, "score")[[c for c in top_cols if c in df.columns]]
    st.dataframe(top, use_container_width=True, hide_index=True)

    _da, _db = st.columns([1, 4])
    with _da:
        if st.button("💬 Ask Job Pal", key="cta_asst_dashboard", type="secondary"):
            st.session_state["_nav_page"] = "Assistant"
            st.rerun()


# ═══════════════════════════════════════════════════════════════════
# REVIEW QUEUE
# ═══════════════════════════════════════════════════════════════════
elif page == "Review Queue":
    page_header("Review Queue", "Jobs worth <em>applying to.</em>")
    rq_col_info, rq_col_thresh, rq_col_view = st.columns([3, 1, 1])
    with rq_col_info:
        st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:12px;color:#8B8B85;margin-bottom:20px">AI-scored matches waiting for your review · status updates save instantly</div>', unsafe_allow_html=True)
    with rq_col_thresh:
        queue_min_score = st.selectbox(
            "Min score",
            [5, 6, 7, 8, 9],
            index=2,  # default 7
            key="rq_min_score",
            help="Lower to see borderline matches. Raise to see only strong fits.",
            label_visibility="collapsed",
        )
    with rq_col_view:
        _rq_view = st.radio(
            "View",
            ["📋 List", "🃏 Cards"],
            key="rq_view_radio",
            horizontal=True,
            label_visibility="collapsed",
        )

    # ── Manual job add ────────────────────────────────────────────────
    with st.expander("➕ Add a job manually", expanded=False):
        st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:#8B8B85;margin-bottom:12px">Found a job outside the scraper? Paste the URL to auto-fill details, or enter them manually.</div>', unsafe_allow_html=True)
        _mj_url = st.text_input("Job URL", key="mj_url", placeholder="https://www.linkedin.com/jobs/view/...")
        _mj_col1, _mj_col2 = st.columns([1, 1])
        with _mj_col1:
            _mj_title    = st.text_input("Job title",   key="mj_title",    placeholder="e.g. Data Scientist")
            _mj_company  = st.text_input("Company",     key="mj_company",  placeholder="e.g. Acme Corp")
            _mj_location = st.text_input("Location",    key="mj_location", placeholder="e.g. Remote / Houston, TX")
        with _mj_col2:
            _mj_score    = st.number_input("Score (1–10)", key="mj_score", min_value=1, max_value=10, value=8,
                                           help="Manually added jobs default to 8 — you picked it, so it matters.")
            _mj_reason   = st.text_input("Why it's interesting (optional)", key="mj_reason",
                                         placeholder="e.g. Strong mission fit, heard from recruiter")
            _mj_wtype    = st.selectbox("Work type", ["Remote", "Hybrid", "Onsite"], key="mj_wtype", index=0)

        _mj_fetch_clicked = st.button("🔍 Fetch details from URL", key="mj_fetch", type="secondary")
        if _mj_fetch_clicked and _mj_url.strip():
            with st.spinner("Fetching job details…"):
                from fetcher import fetch_linkedin_job
                _fetched = fetch_linkedin_job(_mj_url.strip())
            if _fetched.get("company"):
                st.session_state["mj_company"] = _fetched["company"]
            if _fetched.get("description"):
                st.session_state["_mj_fetched_desc"] = _fetched["description"]
                st.success(f"Fetched description ({len(_fetched['description'])} chars).")
            else:
                st.info("Couldn't auto-fill description — fill in manually above and save.")

        if st.button("💾 Add to Review Queue", key="mj_save", type="primary"):
            import hashlib
            _url_clean = _mj_url.strip()
            if not _mj_title.strip() or not _mj_company.strip():
                st.error("Title and company are required.")
            else:
                _mid = hashlib.md5((_url_clean or f"manual-{_mj_title}-{_mj_company}").encode()).hexdigest()[:16]
                if _USER_ID:
                    _mid = hashlib.md5(f"{_USER_ID}-{_url_clean or _mj_title}-{_mj_company}".encode()).hexdigest()[:16]
                _manual_job = {
                    "id":           _mid,
                    "source":       "manual",
                    "title":        _mj_title.strip(),
                    "company":      _mj_company.strip(),
                    "location":     _mj_location.strip() or "Unknown",
                    "url":          _url_clean or "",
                    "description":  st.session_state.pop("_mj_fetched_desc", ""),
                    "status":       "new",
                    "score":        int(_mj_score),
                    "score_reason": _mj_reason.strip() or "Manually added",
                    "work_type":    _mj_wtype,
                    "scored_by":    "manual",
                    "seniority":    "Unknown",
                    "salary_match": "Unknown",
                }
                if _USER_ID:
                    _manual_job["user_id"] = _USER_ID
                from tracker import upsert_jobs
                upsert_jobs([_manual_job])
                _cached_review_queue.clear()
                st.success(f"Added **{_mj_title.strip()} @ {_mj_company.strip()}** to your queue.")
                st.rerun()

    queue = safe_get_queue(min_score=queue_min_score)
    if not queue:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px 24px">
          <h4>Queue is empty</h4>
          <p>All caught up. Run the pipeline to pull in new matches, or lower the min score filter above.</p>
        </div>
        """, unsafe_allow_html=True)
        _qa, _qb = st.columns([1, 4])
        with _qa:
            if st.button("▶ Run Pipeline", type="primary"):
                st.session_state["_nav_page"] = "Run Pipeline"
                st.rerun()
        st.stop()

    # ── Queue filters ─────────────────────────────────────────────
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        hide_old = st.toggle(
            "Hide stale listings (>14 days)",
            value=True,
            help="Postings older than 2 weeks are often filled. Toggle off to see all.",
        )
    with col_f2:
        display_limit = st.selectbox(
            "Show top",
            [25, 50, 100, "All"],
            index=0,
            help="Limit visible jobs — sorted by score before the cutoff.",
        )
    with col_f3:
        all_cats = sorted(set(_categorize(j.get("title","")) for j in queue))
        cat_filter = st.multiselect("Category", all_cats, default=all_cats)

    total_in_db = len(queue)
    unfiltered_queue = queue  # keep reference for purge button — avoids second DB query
    stale_count = sum(1 for j in queue if (_days_in_queue(j) or 0) > 14)
    if hide_old:
        queue = [j for j in queue if (_days_in_queue(j) or 0) <= 14]
    queue = [j for j in queue if _categorize(j.get("title","")) in cat_filter]
    stale_hidden = stale_count if hide_old else 0

    personalize = st.toggle(
        "Learn from my history",
        value=True,
        help="Reorders your queue based on companies, job types, and patterns from roles you've applied to or skipped.",
    )
    weights = {
        "pos_company": 1.2,
        "neg_company": -0.9,
        "good_source": 0.35,
        "title_token_per_hit": 0.2,
        "title_token_cap": 1.0,
        "title_min_hits": 2,
        "neg_token_per_hit": -0.5,
        "neg_token_cap": -1.5,
    }

    if not queue:
        st.info("No jobs match the current filters — try lowering the score threshold or adjusting category filters.")
        st.stop()

    if personalize:
        # Show a plain-English summary of what the system has learned
        ctx = _cached_personalization_ctx(_USER_ID)
        if ctx.get("has_signals"):
            pills = []
            _pos_cos = sorted(ctx.get("pos_companies") or [], key=str.lower)[:4]
            _neg_cos = sorted(ctx.get("neg_companies") or [], key=str.lower)[:4]
            # Only show company signals when they point to different companies
            _pos_cos_clean = [c for c in _pos_cos if c.lower() not in {x.lower() for x in _neg_cos}]
            _neg_cos_clean = [c for c in _neg_cos if c.lower() not in {x.lower() for x in _pos_cos}]
            if _pos_cos_clean:
                pills.append(f"👍 Companies you've pursued: {', '.join(_pos_cos_clean)}")
            if _neg_cos_clean:
                pills.append(f"⚠️ Companies you've passed on: {', '.join(_neg_cos_clean)}")
            if ctx.get("title_tokens"):
                pills.append(f"✅ Role keywords you go for: {', '.join(sorted(ctx['title_tokens'])[:6])}")
            if ctx.get("neg_title_tokens"):
                pills.append(f"⏭️ Role types you tend to skip: {', '.join(sorted(ctx['neg_title_tokens'])[:5])}")
            # Only show expander when there are at least 2 meaningful signal lines
            if len(pills) >= 2:
                with st.expander("What Job Pal has learned about you", expanded=False):
                    for p in pills:
                        st.caption(p)
        else:
            st.caption("Apply to or skip a few jobs and Job Pal will start learning your preferences.")

        rank_queue_with_personalization(queue, weights=weights, ctx=ctx)
    else:
        queue.sort(key=lambda j: j.get("score") or 0, reverse=True)

    display_queue = queue if display_limit == "All" else queue[:int(display_limit)]
    cap_hidden = len(queue) - len(display_queue)

    card_mode = (_rq_view == "🃏 Cards")
    _was_card = st.session_state.get("_rq_card_mode_prev", False)
    if card_mode and not _was_card:
        # Entering card mode: clear passes, bust cache, then rerun so the deck
        # renders from fresh data (not the stale queue already fetched above)
        st.session_state["rq_passed_ids"] = set()
        st.session_state["rq_actioned_ids"] = set()
        _cached_review_queue.clear()
        _cached_personalization_ctx.clear()
        st.session_state["_rq_card_mode_prev"] = card_mode
        st.rerun()
    elif not card_mode and _was_card:
        # Leaving card mode: bust cache so card-actioned jobs don't reappear in list
        st.session_state["rq_actioned_ids"] = set()
        _cached_review_queue.clear()
        st.session_state["_rq_card_mode_prev"] = card_mode
        st.rerun()
    else:
        st.session_state["_rq_card_mode_prev"] = card_mode

    parts = [f"{len(display_queue)} shown"]
    if stale_hidden:
        parts.append(f"{stale_hidden} stale hidden")
    if cap_hidden:
        parts.append(f"{cap_hidden} more below limit")

    summary_col, closed_col, skip_col, purge_col = st.columns([4, 1, 1, 1])
    with summary_col:
        st.markdown(f'<div class="section-label">{" · ".join(parts)}</div>', unsafe_allow_html=True)
    with closed_col:
        if st.button("🔍 Check closed", type="secondary", use_container_width=True,
                     help="Fetch each listing's URL and mark any that say 'no longer accepting applications'"):
            from fetcher import check_closed_listings
            from tracker import mark_closed
            _to_check = [j for j in display_queue if j.get("url")]
            if _to_check:
                with st.spinner(f"Checking {len(_to_check)} listings…"):
                    _closed_ids = check_closed_listings(_to_check)
                if _closed_ids:
                    _n = mark_closed(_closed_ids, user_id=_USER_ID)
                    _cached_review_queue.clear()
                    _cached_personalization_ctx.clear()
                    st.success(f"Marked {_n} listing(s) as closed.")
                    st.rerun()
                else:
                    st.success("All visible listings are still active.")
    with skip_col:
        if st.button(f"Skip all {len(display_queue)}", type="secondary", use_container_width=True,
                     help="Mark every job currently visible as Skipped"):
            for j in display_queue:
                update_status(j["id"], "skipped", user_id=_USER_ID)
                log_event(j["id"], "status_change", "bulk skip", user_id=_USER_ID)
            _cached_review_queue.clear()
            _cached_personalization_ctx.clear()
            st.success(f"Skipped {len(display_queue)} jobs.")
            st.rerun()
    with purge_col:
        stale_jobs = [j for j in unfiltered_queue if (_days_in_queue(j) or 0) > 14]
        if stale_jobs:
            if st.button(f"Dismiss {len(stale_jobs)} stale", type="secondary", help="Mark all jobs older than 14 days as Skipped", use_container_width=True):
                for j in stale_jobs:
                    update_status(j["id"], "skipped", user_id=_USER_ID)
                    log_event(j["id"], "status_change", "stale -> skipped", user_id=_USER_ID)
                _cached_review_queue.clear()
                _cached_personalization_ctx.clear()
                st.success(f"Dismissed {len(stale_jobs)} stale jobs.")
                st.rerun()

    if card_mode:
        _render_swipe_card_mode(display_queue)
    else:
        for job in display_queue:
            job_card(job, "rq", ["applied", "skipped", "rejected"])

    st.markdown("---")
    st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:#8B8B85;margin-bottom:8px">Want to talk through which jobs to apply to?</div>', unsafe_allow_html=True)
    if st.button("💬 Ask Job Pal", key="cta_asst_queue", type="secondary"):
        st.session_state["_nav_page"] = "Assistant"
        st.rerun()


# ═══════════════════════════════════════════════════════════════════
# APPLIED
# ═══════════════════════════════════════════════════════════════════
elif page == "Applied":
    page_header("Applied", "In their <em>inbox.</em>")

    apps = safe_get_apps()
    applied   = [a for a in apps if a.get("status") == "applied"]
    interviews= [a for a in apps if a.get("status") == "interview"]
    rejected  = [a for a in apps if a.get("status") == "rejected"]
    _total_heard = len(interviews) + len(rejected)
    _total_ever  = len(applied) + _total_heard
    _resp_rate   = round(_total_heard / _total_ever * 100) if _total_ever else 0

    _ac1, _ac2, _ac3, _ac4 = st.columns(4)
    metric(_ac1, "Applied",       len(applied))
    metric(_ac2, "Heard Back",    _total_heard, f"{_resp_rate}% response rate")
    metric(_ac3, "Interviews",    len(interviews), accent=bool(interviews))
    metric(_ac4, "Rejected",      len(rejected))

    # ── Dismiss stale applied jobs ────────────────────────────────────
    _stale_applied = [a for a in applied if (_days_in_queue(a) or 0) > 30]
    if _stale_applied:
        _ds_col, _ = st.columns([2, 5])
        with _ds_col:
            if st.button(
                f"Dismiss {len(_stale_applied)} no-response (30+ days)",
                type="secondary",
                use_container_width=True,
                help="Moves applications with no reply after 30 days to 'no response' — keeps history, removes from active view",
            ):
                for _a in _stale_applied:
                    update_status(_a["id"], "no_response", user_id=_USER_ID)
                    log_event(_a["id"], "status_change", "no reply after 30 days", user_id=_USER_ID)
                st.success(f"Moved {len(_stale_applied)} ghost applications to 'no response'.")
                st.rerun()

    if not applied:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px 24px;margin-top:24px">
          <h4>No applications yet</h4>
          <p>Move jobs from Review Queue to Applied once you've submitted them.</p>
        </div>
        """, unsafe_allow_html=True)
        st.stop()

    # ── Gmail rejection scan ──────────────────────────────────────────
    _gmail_on     = st.session_state.get("gmail_scan_enabled", False)
    _gmail_email  = st.session_state.get("gmail_apply_email", "")
    _gmail_pw     = st.session_state.get("gmail_app_password", "")

    # Inline credential prompt — shows when Gmail is enabled but password missing
    if _gmail_on and not _gmail_pw:
        with st.expander("📬 Gmail scanning enabled — enter App Password to activate", expanded=True):
            st.caption("Your password is only kept in memory for this browser session and never stored.")
            _ic1, _ic2 = st.columns([2, 2])
            with _ic1:
                _inline_email = st.text_input(
                    "Gmail address",
                    value=_gmail_email or "tegapeters11@gmail.com",
                    key="inline_gmail_email",
                )
            with _ic2:
                _inline_pw = st.text_input(
                    "App Password (16-char)",
                    type="password",
                    placeholder="xxxx xxxx xxxx xxxx",
                    key="inline_gmail_pw",
                )
            if st.button("Activate scanning", type="primary", key="inline_gmail_save"):
                if _inline_email and _inline_pw:
                    st.session_state["gmail_apply_email"] = _inline_email.strip()
                    st.session_state["gmail_app_password"] = _inline_pw.replace(" ", "").strip()
                    st.rerun()
                else:
                    st.warning("Enter both email and App Password.")

    if _gmail_on and _gmail_email and _gmail_pw:
        # Auto-scan once per session on first Applied page load
        _scan_done = "_gmail_scan_results" in st.session_state
        if not _scan_done:
            with st.spinner("📬 Scanning Gmail for rejections, interviews, and action items…"):
                try:
                    from rejection_scanner import scan_inbox
                    _scan_results = scan_inbox(_gmail_email, _gmail_pw, applied, lookback_days=90)
                    st.session_state["_gmail_scan_results"] = _scan_results
                except ConnectionError as _ce:
                    st.error(str(_ce))
                    st.session_state["_gmail_scan_results"] = None
                except Exception as _ex:
                    st.error(f"Gmail scan failed: {_ex}")
                    st.session_state["_gmail_scan_results"] = None

        _scan_col, _ = st.columns([2, 5])
        with _scan_col:
            if st.button("Re-scan Gmail", type="secondary", use_container_width=True):
                with st.spinner("Scanning inbox…"):
                    try:
                        from rejection_scanner import scan_inbox
                        _scan_results = scan_inbox(_gmail_email, _gmail_pw, applied, lookback_days=90)
                        st.session_state["_gmail_scan_results"] = _scan_results
                    except ConnectionError as _ce:
                        st.error(str(_ce))
                        st.session_state.pop("_gmail_scan_results", None)
                    except Exception as _ex:
                        st.error(f"Scan failed: {_ex}")
                        st.session_state.pop("_gmail_scan_results", None)

        _scan = st.session_state.get("_gmail_scan_results")
        if _scan is not None:
            _n_interview  = len(_scan.get("interview", []))
            _n_action     = len(_scan.get("action", []))
            _n_rejection  = len(_scan.get("rejection", []))
            _n_total      = _n_interview + _n_action + _n_rejection

            if _n_total == 0:
                st.success("Inbox clean — no rejections, interview invites, or action items found. 🎉")
            else:
                # ── Interview invites ──────────────────────────────────
                if _n_interview:
                    st.markdown(f"#### 🟢 Interview Invites ({_n_interview})")
                    _to_interview: list[str] = []
                    for _m in _scan["interview"]:
                        _j = _m["job"]
                        _conf = int(_m["confidence"] * 100)
                        with st.expander(f"{'🟢' if _conf >= 80 else '🟡'} {_j.get('title')} @ {_j.get('company')}  —  {_conf}% match"):
                            st.markdown(f"**From:** {_m['email_from']}")
                            st.markdown(f"**Subject:** {_m['email_subject']}")
                            st.markdown(f"**Date:** {_m['email_date']}")
                            st.caption(_m["snippet"])
                            if st.checkbox("Move to Interview", key=f"int_check_{_j['id']}", value=_conf >= 80):
                                _to_interview.append(_j["id"])
                    if _to_interview:
                        if st.button(f"✅ Mark {len(_to_interview)} as Interview", type="primary", key="btn_mark_interview"):
                            from tracker import update_status
                            for _jid in _to_interview:
                                update_status(_jid, "interview", user_id=_USER_ID)
                            st.success(f"Moved {len(_to_interview)} job(s) to Interview.")
                            st.session_state.pop("_gmail_scan_results", None)
                            st.rerun()

                # ── Action required ────────────────────────────────────
                if _n_action:
                    st.markdown(f"#### 🟡 Action Required ({_n_action})")
                    for _m in _scan["action"]:
                        _j = _m["job"]
                        _conf = int(_m["confidence"] * 100)
                        with st.expander(f"⚡ {_j.get('title')} @ {_j.get('company')}  —  {_conf}% match"):
                            st.markdown(f"**From:** {_m['email_from']}")
                            st.markdown(f"**Subject:** {_m['email_subject']}")
                            st.markdown(f"**Date:** {_m['email_date']}")
                            st.caption(_m["snippet"])
                            st.markdown(f"[Open job posting]({_j.get('url', '#')})")

                # ── Rejections ─────────────────────────────────────────
                if _n_rejection:
                    st.markdown(f"#### 🔴 Rejections ({_n_rejection})")
                    _to_reject: list[str] = []
                    for _m in _scan["rejection"]:
                        _j = _m["job"]
                        _conf = int(_m["confidence"] * 100)
                        with st.expander(f"{'🔴' if _conf >= 80 else '🟡'} {_j.get('title')} @ {_j.get('company')}  —  {_conf}% match"):
                            st.markdown(f"**From:** {_m['email_from']}")
                            st.markdown(f"**Subject:** {_m['email_subject']}")
                            st.markdown(f"**Date:** {_m['email_date']}")
                            st.caption(_m["snippet"])
                            if st.checkbox("Mark as rejected", key=f"rej_check_{_j['id']}", value=_conf >= 80):
                                _to_reject.append(_j["id"])
                    if _to_reject:
                        if st.button(f"✅ Mark {len(_to_reject)} as Rejected", type="primary", key="btn_mark_rejected"):
                            from tracker import update_status
                            for _jid in _to_reject:
                                update_status(_jid, "rejected", user_id=_USER_ID)
                            st.success(f"Marked {len(_to_reject)} job(s) as rejected.")
                            st.session_state.pop("_gmail_scan_results", None)
                            st.rerun()

    elif _gmail_on and (not _gmail_email or not _gmail_pw):
        st.caption("📬 Gmail scanning is enabled — add your credentials on the Setup page to activate it.")

    # ── Filters ──────────────────────────────────────────────────────
    f1, f2, f3 = st.columns([3, 2, 1])
    with f1:
        search = st.text_input("Search", placeholder="Company or title...")
    with f2:
        _date_options = {"Last 30 days": 30, "Last 60 days": 60, "Last 90 days": 90, "All time": 0}
        _date_label   = st.selectbox("Date range", list(_date_options.keys()), index=0)
        _date_days    = _date_options[_date_label]
    with f3:
        cover_letter_only = st.toggle("Cover letter only", value=False)

    if search:
        applied = [a for a in applied if
                   search.lower() in (a.get("title") or "").lower() or
                   search.lower() in (a.get("company") or "").lower()]
    if _date_days:
        from datetime import datetime, timezone, timedelta
        _cutoff = datetime.now(timezone.utc) - timedelta(days=_date_days)
        def _after_cutoff(job):
            _ts = job.get("created_at") or ""
            if not _ts:
                return True
            try:
                _dt = datetime.fromisoformat(_ts.replace("Z", "+00:00"))
                return _dt >= _cutoff
            except Exception:
                return True
        applied = [a for a in applied if _after_cutoff(a)]
    if cover_letter_only:
        applied = [a for a in applied if a.get("cover_letter")]

    st.markdown(f'<div class="section-label">{len(applied)} applications</div>', unsafe_allow_html=True)

    for job in sorted(applied, key=lambda x: x.get("score") or 0, reverse=True):
        job_card(job, "ap", ["interview", "rejected", "skipped", "application_closed"])


# ═══════════════════════════════════════════════════════════════════
# INTERVIEWS
# ═══════════════════════════════════════════════════════════════════
elif page == "Interviews":
    page_header("Interviews", "You're in the <em>room.</em>")

    apps = safe_get_apps()
    interviews = [a for a in apps if a.get("status") == "interview"]

    if not interviews:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px 24px">
          <h4>No active interviews</h4>
          <p>When you land a phone screen or interview, move it here from Applied to track prep and outcomes.</p>
        </div>
        """, unsafe_allow_html=True)
        _ia, _ib = st.columns([1,4])
        with _ia:
            if st.button("✅ Go to Applied", type="secondary"):
                st.session_state["_nav_page"] = "Applied"
                st.rerun()
        st.stop()

    _iv1, _iv2, _iv3 = st.columns(3)
    metric(_iv1, "Active",    len(interviews))
    metric(_iv2, "Avg Score", f"{sum((j.get('score') or 0) for j in interviews)/len(interviews):.1f}", "ai score", accent=True)
    _days_list = [_days_in_queue(j) for j in interviews if _days_in_queue(j) is not None]
    metric(_iv3, "Avg Days",  f"{round(sum(_days_list)/len(_days_list))}" if _days_list else "—", "in pipeline")

    st.markdown(f'<div class="section-label" style="margin-top:24px">{len(interviews)} active interview{"s" if len(interviews) != 1 else ""}</div>', unsafe_allow_html=True)

    _prep_a, _prep_b = st.columns([4, 1])
    with _prep_b:
        if st.button("🤖 Prep with AI", type="secondary", use_container_width=True):
            st.session_state["_nav_page"] = "Assistant"
            st.rerun()

    for job in sorted(interviews, key=lambda x: x.get("score") or 0, reverse=True):
        job_card(job, "iv", ["rejected", "skipped", "application_closed"], expanded=True)


# ═══════════════════════════════════════════════════════════════════
# ALL APPLICATIONS
# ═══════════════════════════════════════════════════════════════════
elif page == "All Applications":
    page_header("All Applications", "Full <em>history.</em>")

    apps = safe_get_apps()
    if not apps:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px 24px">
          <h4>No history yet</h4>
          <p>Run the pipeline to start tracking jobs.</p>
        </div>
        """, unsafe_allow_html=True)
        st.stop()

    df = pd.DataFrame(apps)

    # ── Status summary bar ────────────────────────────────────────
    _sc = df["status"].value_counts().to_dict()
    _chip_styles = {
        "new":                "background:#1a1f0a;color:#D4FF3A;border:1px solid #2a3510",
        "applied":            "background:#0a1a1f;color:#3ad4ff;border:1px solid #102a35",
        "interview":          "background:#1a0f1a;color:#d43aff;border:1px solid #2a1035",
        "rejected":           "background:#1f0a0a;color:#ff3a3a;border:1px solid #350f0f",
        "skipped":            "background:#1a1a1a;color:#666;border:1px solid #333",
        "application_closed": "background:#1a1208;color:#ff9a3a;border:1px solid #352010",
    }
    _chips = " ".join(
        f'<span class="sb-chip" style="{_chip_styles.get(s, "background:#1a1a1a;color:#888")}">'
        f'{s.replace("application_closed","closed")} · {n}</span>'
        for s, n in sorted(_sc.items(), key=lambda x: -x[1])
    )
    st.markdown(f'<div class="status-bar">{_chips}</div>', unsafe_allow_html=True)

    df["category"] = df["title"].apply(_categorize)
    categories = sorted(df["category"].unique().tolist())

    col1, col2 = st.columns(2)
    with col1:
        status_filter = st.multiselect(
            "Status", df["status"].unique().tolist(),
            default=df["status"].unique().tolist(),
        )
    with col2:
        category_filter = st.multiselect(
            "Category", categories, default=categories,
        )

    col3, col4 = st.columns(2)
    with col3:
        min_score = st.slider("Min score", 0, 10, 0)
    with col4:
        search = st.text_input("Search title / company")

    filtered = df[
        df["status"].isin(status_filter) &
        df["category"].isin(category_filter) &
        (df["score"].fillna(0) >= min_score)
    ]
    if search:
        mask = (
            filtered["title"].str.contains(search, case=False, na=False) |
            filtered["company"].str.contains(search, case=False, na=False)
        )
        filtered = filtered[mask]

    st.markdown(f'<div class="section-label">{len(filtered)} results</div>', unsafe_allow_html=True)

    display_cols = ["category", "title", "company", "location", "score", "status", "seniority", "salary_match", "source"]
    available = [c for c in display_cols if c in filtered.columns]
    st.dataframe(filtered[available], use_container_width=True, hide_index=True)

    st.markdown('<div class="section-label" style="margin-top:28px">Quick Status Update</div>', unsafe_allow_html=True)
    job_options = {
        f"{r['title']} @ {r.get('company','?')} (score: {r.get('score','?')})": r["id"]
        for _, r in filtered.iterrows()
    }
    if job_options:
        selected_label = st.selectbox("Select job", list(job_options.keys()))
        selected_id = job_options[selected_label]
        new_status = st.selectbox("New status", ["new", "applied", "interview", "rejected", "skipped", "application_closed"])
        if st.button("Update Status", type="primary"):
            update_status(selected_id, new_status, user_id=_USER_ID)
            _evt_map = {"applied": "applied", "interview": "interview_scheduled", "rejected": "rejected"}
            if new_status in _evt_map:
                log_event(selected_id, _evt_map[new_status], f"status moved to {new_status}", user_id=_USER_ID)
            else:
                log_event(selected_id, "status_change", f"-> {new_status}", user_id=_USER_ID)
            _cached_review_queue.clear()
            _cached_personalization_ctx.clear()
            st.success(f"Updated → {new_status}")
            st.rerun()

    st.markdown("---")
    st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:#8B8B85;margin-bottom:8px">Need help with a follow-up or status update?</div>', unsafe_allow_html=True)
    if st.button("💬 Ask Job Pal", key="cta_asst_all_apps", type="secondary"):
        st.session_state["_nav_page"] = "Assistant"
        st.rerun()


# ═══════════════════════════════════════════════════════════════════
# RUN PIPELINE
# ═══════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════
# EVENTS
# ═══════════════════════════════════════════════════════════════════
elif page == "Events":
    page_header("Networking", "Events near <em>you.</em>")

    from agent import score_events, extract_cities_from_resume

    # ── City list: detected from resume + manually added ─────────
    resume_text = st.session_state.get("resume_text")
    resume_cities = extract_cities_from_resume(resume_text) if resume_text else []
    if "Houston" not in resume_cities:
        resume_cities = ["Houston"] + resume_cities

    manual_cities = st.session_state.get("ev_manual_cities", [])
    all_cities = resume_cities + [c for c in manual_cities if c not in resume_cities]

    # ── Controls ──────────────────────────────────────────────────
    ctrl_col1, ctrl_col2 = st.columns([3, 1])
    with ctrl_col1:
        ev_min_score = st.selectbox(
            "Min relevance",
            [1, 3, 5, 7],
            index=2,
            label_visibility="collapsed",
            help="Minimum relevance score to show",
        )
    with ctrl_col2:
        refresh_btn = st.button("Refresh Events", type="primary", use_container_width=True)

    # City tags — auto-detected + editable
    st.markdown('<div class="section-label">Cities</div>', unsafe_allow_html=True)
    selected_cities = st.multiselect(
        "Cities to scrape",
        options=all_cities,
        default=all_cities,
        label_visibility="collapsed",
        help="Uncheck a city to exclude it. Use the field below to add more.",
        key="ev_city_select",
    )

    add_col, btn_col = st.columns([4, 1])
    with add_col:
        new_city = st.text_input(
            "Add another city",
            placeholder="e.g. Dallas, Austin, New York…",
            key="ev_new_city",
            label_visibility="collapsed",
        )
    with btn_col:
        if st.button("＋ Add city", key="ev_add_city", use_container_width=True):
            city_to_add = new_city.strip()
            if city_to_add and city_to_add not in all_cities:
                st.session_state.setdefault("ev_manual_cities", []).append(city_to_add)
                st.rerun()
            elif not city_to_add:
                st.toast("Enter a city name first.")
            else:
                st.toast(f"{city_to_add} is already in the list.")

    cities_to_scrape = selected_cities
    if not cities_to_scrape:
        st.warning("Select at least one city to search.")
        st.stop()

    if refresh_btn:
        with st.spinner(f"Scraping events in {', '.join(cities_to_scrape)}..."):
            from scrapers.events import scrape_events
            raw = scrape_events(cities=cities_to_scrape)

        if not raw:
            st.warning("No events found — existing events kept. Try again or add more cities.")
        else:
            st.info(f"Found {len(raw)} events across {len(cities_to_scrape)} cities — scoring against your resume...")
            scored = score_events(raw, resume_text=resume_text)

            with st.spinner("Saving to your account..."):
                try:
                    # Delete AFTER a successful scrape so a failed run never wipes the DB
                    delete_all_events(user_id=_USER_ID)
                    upsert_events(scored, user_id=_USER_ID)
                except Exception as e:
                    st.error(
                        f"Could not save events ({e}). "
                        "Check that EVENTS_SERVICE_ROLE_KEY is set in Streamlit Cloud secrets."
                    )
                    st.stop()

            st.success(f"✓ {len(scored)} events loaded. Showing {ev_min_score}+ relevance below.")
            st.rerun()

    # ── Load saved events ─────────────────────────────────────────
    # Quietly prune past events on every page load
    try:
        from tracker import delete_past_events
        delete_past_events(user_id=_USER_ID)
    except Exception:
        pass

    try:
        events = get_events(user_id=_USER_ID, min_score=ev_min_score,
                            status_filter=["new", "interested", "attending", "skipped"])
    except RuntimeError:
        st.warning(
            "Events database not configured yet. "
            "Add `EVENTS_SUPABASE_URL` and `EVENTS_SUPABASE_KEY` to your "
            "Streamlit Cloud secrets (Settings → Secrets), then reboot the app."
        )
        st.stop()

    if not events:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px">
          <h4>No events loaded yet</h4>
          <p>Enter your city above and click <b>Refresh Events</b> to scrape local networking events
          scored against your resume.</p>
        </div>
        """, unsafe_allow_html=True)
        st.stop()

    # ── Filters ───────────────────────────────────────────────────
    all_sources = sorted(set(e.get("source","") for e in events))
    f1, f2, f3 = st.columns([2, 2, 2])
    with f1:
        src_filter = st.multiselect("Source", all_sources, default=all_sources, key="ev_src")
    with f2:
        ev_search = st.text_input("Search", placeholder="AI, data, Python…", key="ev_search",
                                  label_visibility="collapsed")
    with f3:
        _ev_sort = st.selectbox(
            "Sort",
            ["Relevance", "Date (soonest first)"],
            index=0,
            key="ev_sort",
            label_visibility="collapsed",
        )

    if src_filter:
        events = [e for e in events if e.get("source") in src_filter]
    if ev_search:
        q = ev_search.lower()
        events = [e for e in events if q in (e.get("title","") + e.get("description","") +
                                              e.get("organizer","")).lower()]

    # ── Status tabs ──────────────────────────────────────────────
    _n_all        = len([e for e in events if e.get("status") != "skipped"])
    _n_interested = len([e for e in events if e.get("status") == "interested"])
    _n_attending  = len([e for e in events if e.get("status") == "attending"])
    _status_tab, _interested_tab, _attending_tab = st.tabs([
        f"All  ({_n_all})",
        f"Interested  ({_n_interested})",
        f"Attending  ({_n_attending})",
    ])

    def _render_events(ev_list, tab_key):
        if not ev_list:
            st.markdown(
                '<div class="pipeline-card" style="text-align:center;padding:24px">'
                '<p>No events in this filter.</p></div>',
                unsafe_allow_html=True,
            )
            return
        _sort_mode = st.session_state.get("ev_sort", "Relevance")
        if _sort_mode == "Date (soonest first)":
            from datetime import datetime, timezone
            def _parse_date(e):
                raw = e.get("start_date") or ""
                try:
                    dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
                    return dt.replace(tzinfo=timezone.utc) if dt.tzinfo is None else dt
                except Exception:
                    return datetime.max.replace(tzinfo=timezone.utc)
            ev_list = sorted(ev_list, key=_parse_date)
            _sort_label = "date"
        else:
            ev_list = sorted(ev_list, key=lambda e: (-(e.get("relevance_score") or 0), e.get("start_date") or ""))
            _sort_label = "relevance"
        st.markdown(f'<div class="section-label">{len(ev_list)} events · sorted by {_sort_label}</div>',
                    unsafe_allow_html=True)
        STATUS_COLORS = {"new": "#4A4A45", "interested": "#f5c518", "attending": "#D4FF3A"}
        for ev in ev_list:
            score  = ev.get("relevance_score") or 0
            status = ev.get("status", "new")
            status_dot = f'<span style="color:{STATUS_COLORS.get(status,"#4A4A45")};font-size:10px">● {status.upper()}</span>'
            raw_date = ev.get("start_date", "")
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
                date_str = dt.strftime("%a %b %-d · %-I:%M %p")
            except Exception:
                date_str = raw_date[:16] if raw_date else "Date TBD"
            eid = ev["id"]
            with st.expander(f"{score}/10  —  {ev.get('title','')[:60]}", expanded=False,
                             key=f"ev_exp_{tab_key}_{eid}"):
                top_col, btn_col = st.columns([3, 1])
                with top_col:
                    st.markdown(
                        f'<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;'
                        f'color:#8B8B85;margin-bottom:8px">'
                        f'{date_str} &nbsp;·&nbsp; '
                        f'{ev.get("location","")[:40]} &nbsp;·&nbsp; '
                        f'{ev.get("organizer","")[:35]}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                    if ev.get("relevance_reason"):
                        st.caption(f"Relevance: {ev['relevance_reason']}")
                    if ev.get("description"):
                        st.markdown(
                            f'<div class="cover-letter" style="max-height:140px;overflow-y:auto;'
                            f'font-size:11px">{ev["description"][:600]}</div>',
                            unsafe_allow_html=True,
                        )
                    if ev.get("url"):
                        st.markdown(f"[View event ↗]({ev['url']})")
                with btn_col:
                    st.markdown(f'<div style="margin-bottom:8px">{status_dot}</div>',
                                unsafe_allow_html=True)
                    if st.button("Interested", key=f"ev_int_{tab_key}_{eid}", use_container_width=True):
                        update_event_status(eid, "interested", user_id=_USER_ID)
                        st.rerun()
                    if st.button("Attending", key=f"ev_att_{tab_key}_{eid}", use_container_width=True,
                                 type="primary"):
                        update_event_status(eid, "attending", user_id=_USER_ID)
                        st.rerun()
                    if st.button("Skip", key=f"ev_skip_{tab_key}_{eid}", use_container_width=True):
                        update_event_status(eid, "skipped", user_id=_USER_ID)
                        st.rerun()

    with _status_tab:
        _render_events([e for e in events if e.get("status") != "skipped"], "all")
    with _interested_tab:
        _render_events([e for e in events if e.get("status") == "interested"], "int")
    with _attending_tab:
        _render_events([e for e in events if e.get("status") == "attending"], "att")

    st.markdown("---")
    st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:#8B8B85;margin-bottom:8px">Want networking tips or event prep help?</div>', unsafe_allow_html=True)
    if st.button("Ask Job Pal", key="cta_asst_events", type="secondary"):
        st.session_state["_nav_page"] = "Assistant"
        st.rerun()


elif page == "Run Pipeline":
    page_header("Run Pipeline", "Scrape. Score. <em>Apply.</em>")

    # ── Resume gate ──────────────────────────────────────────────
    resume_text = st.session_state.get("resume_text")
    if not resume_text:
        st.markdown("""
        <div class="pipeline-card" style="text-align:center;padding:40px 24px">
          <h4>No resume loaded</h4>
          <p>Go to Setup and paste or upload your resume first — the pipeline uses it to score every job for fit.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("⚙ Go to Setup", type="primary"):
            st.session_state["_nav_page"] = "Setup"
            st.rerun()
        st.stop()

    # ── Pre-flight checklist ──────────────────────────────────────
    _roles   = st.session_state.get("target_roles") or []
    _salary  = st.session_state.get("min_salary", 0)
    _pref_locs = st.session_state.get("preferred_locations") or []
    _resume_kb = round(len(resume_text) / 1024, 1)
    _salary_label = f"${_salary:,}+" if _salary else "No minimum set"
    _roles_label  = ", ".join(_roles[:3]) + ("…" if len(_roles) > 3 else "") if _roles else None
    _locs_label   = ", ".join(_pref_locs[:3]) + ("…" if len(_pref_locs) > 3 else "") if _pref_locs else "All locations (set in Setup)"

    st.markdown(f"""
    <div class="preflight">
      <div class="pf-row"><span class="ok">✓</span> Resume loaded &nbsp;<span style="color:#4A4A45">({_resume_kb} KB)</span></div>
      <div class="pf-row">
        {'<span class="ok">✓</span>' if _roles else '<span class="warn">✗</span>'}
        {'Target roles: <span style="color:#F5F4EE">' + _roles_label + '</span>' if _roles else 'No target roles set — go to Setup before running'}
      </div>
      <div class="pf-row"><span class="ok">✓</span> Salary floor: <span style="color:#F5F4EE">{_salary_label}</span></div>
      <div class="pf-row"><span class="ok">✓</span> Locations: <span style="color:#F5F4EE">{_locs_label}</span></div>
      <div class="pf-row"><span class="ok">✓</span> Sources: LinkedIn · Adzuna · USAJobs · The Muse · RemoteOK · Remotive · We Work Remotely · Jobicy</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Main controls ─────────────────────────────────────────────
    _mc1, _mc2 = st.columns([2, 1])
    with _mc1:
        scoring_mode = st.selectbox(
            "Scoring mode",
            ["hybrid", "cheap"],
            index=0,
            help="hybrid = keyword pre-filter + Claude (recommended) · cheap = heuristic only, no LLM",
        )
    with _mc2:
        enable_letters = st.checkbox(
            "Generate cover letters",
            value=(scoring_mode != "cheap"),
            help="Writes a tailored cover letter for every job scoring 8+",
        )

    _mode_info = {
        "claude": "🔵 Claude scores every job — most accurate. Recommended.",
        "hybrid": "🟡 Keyword pre-filter, then Claude for anything above the threshold. ~60% of Claude cost.",
        "cheap":  "⚪ Keyword heuristic only. No LLM, no API calls. Dev/testing only.",
    }
    st.caption(_mode_info.get(scoring_mode, ""))

    # ── Advanced options ──────────────────────────────────────────
    with st.expander("⚙ Advanced options"):
        _adv1, _adv2 = st.columns(2)
        with _adv1:
            run_label = st.text_input("Run label", value="", placeholder="e.g. embed-jul15")
            hybrid_min = st.slider("Hybrid Claude threshold", 1, 10, 6, disabled=(scoring_mode != "hybrid"),
                                   help="Jobs at or above this cheap-score are escalated to Claude.")
        with _adv2:
            run_note = st.text_area("Run note", value="", height=100,
                                    placeholder="Notes on this run (query changes, observations).")

    if st.button("▶ Run Pipeline", type="primary", use_container_width=True):
        _pipeline_roles = st.session_state.get("target_roles") or []
        if not _pipeline_roles:
            st.error("No target roles set — go to **Setup** and add the roles you're looking for before running.")
            st.stop()

        # Auto-prune unreviewed jobs older than 14 days before each run
        from tracker import prune_stale_queue
        _pruned = prune_stale_queue(user_id=_USER_ID, days=14)
        if _pruned:
            st.info(f"🗑️ Cleared {_pruned} unreviewed jobs older than 14 days — fresh listings only.")

        # Closed-listing sweep: check 6-and-under jobs sitting 5+ days
        from tracker import get_stale_listings, mark_closed
        _stale = get_stale_listings(user_id=_USER_ID, min_days=5, max_score=6)
        if _stale:
            with st.spinner(f"Checking {len(_stale)} stale listings for closed postings…"):
                from fetcher import check_closed_listings
                _closed_ids = check_closed_listings(_stale)
                if _closed_ids:
                    _n_closed = mark_closed(_closed_ids, user_id=_USER_ID)
                    st.info(f"🚫 Marked {_n_closed} listing(s) as closed — removed from your queue.")

        _vertical = st.session_state.get("vertical", "tech")
        _schedule_pref = st.session_state.get("schedule_pref") or None
        _src_label = "higher-ed faculty boards" if _vertical == "academic" else "8 sources"
        with st.spinner(f"Scraping jobs from {_src_label} in parallel…"):
            from scrapers import scrape_all
            jobs = scrape_all(
                target_roles=_pipeline_roles,
                min_salary=st.session_state.get("min_salary", 0),
                locations=st.session_state.get("preferred_locations") or [],
                vertical=_vertical,
            )

        st.info(f"Scraped {len(jobs)} jobs total")

        from tracker import get_seen_ids, get_actioned_fingerprints
        seen = get_seen_ids(user_id=_USER_ID)
        actioned = get_actioned_fingerprints(user_id=_USER_ID)
        new_jobs = [j for j in jobs if _scope_id(j["id"], _USER_ID) not in seen]

        # Cross-source dedup: drop jobs whose (title, company) already has a
        # non-new status (skipped/applied/etc) regardless of which source/URL
        # they came from this time.
        new_jobs = [
            j for j in new_jobs
            if ((j.get("title") or "").lower().strip(),
                (j.get("company") or "").lower().strip()) not in actioned
        ]

        _seen_tc: dict[tuple, dict] = {}
        for j in new_jobs:
            key = (j.get("title", "").strip().lower(), j.get("company", "").strip().lower())
            existing = _seen_tc.get(key)
            if existing is None or len(j.get("description") or "") > len(existing.get("description") or ""):
                _seen_tc[key] = j
        new_jobs = list(_seen_tc.values())

        if len(new_jobs) > BETA_JOB_LIMIT:
            st.info(f"Beta limit: pre-ranking {len(new_jobs)} new jobs, keeping top {BETA_JOB_LIMIT}…")
            from agent import score_job_cheap
            _resume = st.session_state.get("resume_text") or ""
            _roles  = st.session_state.get("target_roles")
            _sal    = st.session_state.get("min_salary", 0)
            for j in new_jobs:
                score_job_cheap(j, resume_text=_resume, target_roles=_roles, min_salary=_sal, vertical=_vertical)
            new_jobs = sorted(new_jobs, key=lambda j: j.get("score") or 0, reverse=True)[:BETA_JOB_LIMIT]
            for j in new_jobs:
                j["score"] = None
        else:
            st.info(f"{len(new_jobs)} new (unseen) jobs to score")

        li_empty = [j for j in new_jobs if "linkedin.com" in (j.get("url") or "") and not j.get("description")]
        if li_empty:
            with st.spinner(f"Fetching full descriptions for {len(li_empty)} LinkedIn jobs…"):
                from fetcher import enrich_jobs
                new_jobs = enrich_jobs(new_jobs)

        if not new_jobs:
            st.success("Queue is up to date — no new jobs to score.")
        else:
            progress = st.progress(0, text="Starting AI scoring…")
            progress.progress(10, text=f"Scoring {len(new_jobs)} jobs against your resume…")
            from agent import process_jobs
            all_scored, qualified, stage_timings = process_jobs(
                new_jobs,
                verbose=False,
                resume_text=resume_text,
                scoring_backend=scoring_mode,
                enable_cover_letters=enable_letters,
                hybrid_claude_min_score=hybrid_min,
                min_salary=st.session_state.get("min_salary", 0),
                target_roles=st.session_state.get("target_roles") or [],
                vertical=_vertical,
                schedule_pref=_schedule_pref,
            )

            progress.progress(80, text=f"Saving {len(qualified)} qualified jobs…")
            from tracker import upsert_jobs
            import time as _time
            _t_save = _time.perf_counter()
            upsert_jobs(all_scored, user_id=_USER_ID)
            stage_timings["save_s"] = round(_time.perf_counter() - _t_save, 1)
            total_s = sum(v for v in stage_timings.values())
            stage_timings["total_s"] = round(total_s, 1)

            log_experiment_run(
                run_label=run_label.strip(),
                scoring_mode=scoring_mode,
                hybrid_threshold=hybrid_min,
                cover_letters_enabled=enable_letters,
                jobs_scraped=len(jobs),
                jobs_new=len(new_jobs),
                jobs_qualified=len(qualified),
                note=run_note.strip(),
                user_id=_USER_ID,
                total_seconds=stage_timings["total_s"],
                timing_json=stage_timings,
            )

            progress.progress(100, text="Done.")
            _timing_str = f"⏱ {stage_timings['total_s']:.0f}s  (score {stage_timings.get('pass2_s',0):.0f}s · enrich {stage_timings.get('enrich_s',0):.0f}s)"
            st.success(f"✓ Pipeline complete — **{len(qualified)} jobs scored {REVIEW_MIN_SCORE}+** · {_timing_str}")

            if qualified:
                st.markdown('<div class="section-label" style="margin-top:20px">New qualified jobs</div>', unsafe_allow_html=True)
                q_df = pd.DataFrame(qualified).reindex(columns=["title", "company", "score", "scored_by", "score_reason", "seniority"])
                st.dataframe(q_df, use_container_width=True, hide_index=True)
                if st.button("📋 Go to Review Queue", type="primary"):
                    _cached_review_queue.clear()
                    _cached_personalization_ctx.clear()
                    st.session_state["_nav_page"] = "Review Queue"
                    st.rerun()

    # ── Scraper health ───────────────────────────────────────────────
    st.markdown('<div class="section-label" style="margin-top:36px">Scraper Health</div>', unsafe_allow_html=True)
    try:
        from tracker import get_source_freshness
        freshness = get_source_freshness(user_id=_USER_ID)
        if freshness:
            _STATUS_COLOR = {"fresh": "#D4FF3A", "ok": "#8BC34A", "stale": "#FF6B6B", "unknown": "#888"}
            _STATUS_LABEL = {"fresh": "Fresh", "ok": "OK", "stale": "Stale", "unknown": "?"}
            cols = st.columns(len(freshness))
            for col, f in zip(cols, freshness):
                status = f["status"]
                color  = _STATUS_COLOR.get(status, "#888")
                label  = _STATUS_LABEL.get(status, "?")
                age    = f["age_hours"]
                if age is None:
                    age_str = "—"
                elif age < 1:
                    age_str = "<1h ago"
                elif age < 24:
                    age_str = f"{age:.0f}h ago"
                else:
                    age_str = f"{age/24:.1f}d ago"
                col.markdown(
                    f'<div style="background:#1C1C18;border:1px solid #2A2A25;border-radius:8px;padding:12px 14px">'
                    f'<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#888;text-transform:uppercase;letter-spacing:0.1em">{f["source"]}</div>'
                    f'<div style="font-size:18px;font-weight:600;color:#F5F4EE;margin:4px 0">{f["count"]}</div>'
                    f'<div style="font-size:11px;color:#888;margin-bottom:6px">{age_str}</div>'
                    f'<span style="background:{color}20;color:{color};font-size:10px;font-family:\'JetBrains Mono\',monospace;'
                    f'padding:2px 8px;border-radius:4px;text-transform:uppercase">{label}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            stale_sources = [f["source"] for f in freshness if f["status"] == "stale"]
            if stale_sources:
                st.warning(f"⚠️ Stale data from: {', '.join(stale_sources)} — run the pipeline to refresh.")
        else:
            st.caption("No scrape data yet.")
    except Exception as e:
        st.caption(f"Health check unavailable ({e})")

    st.markdown('<div class="section-label" style="margin-top:36px">Current Stats</div>', unsafe_allow_html=True)
    try:
        apps = safe_get_apps()
        if apps:
            df = pd.DataFrame(apps)
            c1, c2, c3 = st.columns(3)
            metric(c1, "Total Tracked", len(df))
            metric(c2, "In Queue",      len(df[df["status"] == "new"]), "awaiting review")
            metric(c3, "Applied",       len(df[df["status"] == "applied"]), accent=True)
        else:
            st.info("No data yet — run the pipeline to get started.")
    except Exception as e:
        st.warning(f"Could not load stats — check Supabase connection. ({e})")

    recent_runs = get_recent_runs(limit=8, user_id=_USER_ID)
    if recent_runs:
        st.markdown('<div class="section-label" style="margin-top:24px">Recent Experiment Runs</div>', unsafe_allow_html=True)
        run_df = enrich_experiment_runs_df(pd.DataFrame(recent_runs))
        keep_cols = [
            "created_at", "run_label", "cost_mode", "scoring_mode", "hybrid_threshold",
            "cover_letters_enabled", "jobs_scraped", "jobs_new", "jobs_qualified",
            "qualified_pct", "total_seconds", "note",
        ]
        cols = [c for c in keep_cols if c in run_df.columns]
        display_df = run_df[cols].copy()
        if "qualified_pct" in display_df.columns:
            display_df["qualified_pct"] = display_df["qualified_pct"].apply(
                lambda x: f"{x}%" if x is not None and pd.notna(x) else "—"
            )
        with_sub = run_df[run_df["jobs_new"].fillna(0) > 0]
        if not with_sub.empty and with_sub["qualified_pct"].notna().any():
            best = with_sub.loc[with_sub["qualified_pct"].idxmax()]
            lbl = best["run_label"] if pd.notna(best["run_label"]) else "(no label)"
            st.caption(
                f"Best qualified rate in this list: **{lbl}** — "
                f"{best['qualified_pct']}% qualified "
                f"({int(best['jobs_qualified'] or 0)}/{int(best['jobs_new'] or 0)} new jobs)"
            )
        st.dataframe(display_df, use_container_width=True, hide_index=True)
        csv_buf = io.StringIO()
        run_df.to_csv(csv_buf, index=False)
        st.download_button(
            "Download recent runs (CSV)",
            csv_buf.getvalue(),
            file_name="pipeline_runs_recent.csv",
            mime="text/csv",
            key="download_pipeline_runs_csv",
        )

    # ── Clear queue (pipeline-only reset, keeps resume) ───────────
    st.markdown('<div class="section-label" style="margin-top:36px"></div>', unsafe_allow_html=True)
    cq_col, _ = st.columns([1, 3])
    with cq_col:
        if st.button("Clear Review Queue", key="btn_clear_queue", use_container_width=True,
                     help="Deletes all unreviewed jobs so you can run a fresh pipeline. Keeps applied/interview history."):
            if st.session_state.get("_confirm_clear_queue"):
                clear_queue(user_id=_USER_ID)
                st.session_state.pop("_confirm_clear_queue", None)
                st.success("✓ Review queue cleared. Run Pipeline to refill it.")
                st.rerun()
            else:
                st.session_state["_confirm_clear_queue"] = True
                st.warning("Clears all unreviewed jobs. Click again to confirm.")

# ── Assistant ───────────────────────────────────────────────────────
elif page == "Assistant":
    page_header("Assistant", "Your job search <em>co-pilot.</em>")

    resume_text = st.session_state.get("resume_text") or ""
    if not resume_text:
        st.warning("No resume loaded. Go to **Setup** first — the assistant uses it to personalise answers.")
        st.stop()

    import anthropic as _anthropic

    # ── Load chat history once per session ─────────────────────────
    if "_chat_history_loaded" not in st.session_state:
        st.session_state["_chat_history"] = (
            load_chat_history(_USER_ID) if _USER_ID else []
        )
        st.session_state["_chat_history_loaded"] = True
    if "_chat_history" not in st.session_state:
        st.session_state["_chat_history"] = []

    # ── Pull live pipeline data (single query) ──────────────────────
    try:
        _asst_apps = get_all_applications(user_id=_USER_ID) or []
    except Exception:
        _asst_apps = []

    _asst_new        = [a for a in _asst_apps if a.get("status") == "new"]
    _asst_applied    = [a for a in _asst_apps if a.get("status") == "applied"]
    _asst_interviews = [a for a in _asst_apps if a.get("status") == "interview"]
    _asst_strong     = len([j for j in _asst_new if (j.get("score") or 0) >= 8])
    _asst_solid      = len([j for j in _asst_new if (j.get("score") or 0) == 7])
    _asst_consider   = len([j for j in _asst_new if (j.get("score") or 0) == 6])
    _asst_queue      = sorted(
        [a for a in _asst_new if (a.get("score") or 0) >= 6],
        key=lambda x: x.get("score") or 0, reverse=True
    )

    def _save_history():
        if _USER_ID:
            try:
                save_chat_history(_USER_ID, st.session_state["_chat_history"])
            except Exception:
                pass

    def _stream_response(history: list[dict]) -> str:
        _client = _anthropic.Anthropic()
        _sys = _build_assistant_system_prompt("Assistant")

        def _gen():
            with _client.messages.stream(
                model="claude-sonnet-4-6",
                max_tokens=1500,
                system=[{"type": "text", "text": _sys, "cache_control": {"type": "ephemeral"}}],
                messages=history,
                extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
            ) as stream:
                for text in stream.text_stream:
                    yield text

        with st.chat_message("assistant", avatar="🤖"):
            response = st.write_stream(_gen())
        return response

    # ── Layout: chat (65%) | slim context panel (35%) ──────────────
    chat_col, ctx_col = st.columns([13, 7])

    # ── RIGHT PANEL — pipeline snapshot + job selector only ─────────
    with ctx_col:
        # Pipeline snapshot
        st.markdown(f"""
        <div style="background:#131315;border:1px solid #1f1f22;border-radius:8px;padding:14px 16px;margin-bottom:12px">
          <div style="font-family:'JetBrains Mono',monospace;font-size:10px;color:#4A4A45;letter-spacing:0.1em;margin-bottom:10px">PIPELINE</div>
          <div style="font-family:'JetBrains Mono',monospace;font-size:12px;color:#F5F4EE;line-height:2">
            <span style="color:#D4FF3A;font-size:15px">{_asst_strong}</span> <span style="color:#4A4A45">strong (8+)</span><br>
            <span style="color:#F5C842;font-size:15px">{_asst_solid}</span> <span style="color:#4A4A45">solid (7)</span><br>
            <span style="color:#8B8B85;font-size:15px">{_asst_consider}</span> <span style="color:#4A4A45">consider (6)</span>
          </div>
          <div style="border-top:1px solid #1f1f22;margin-top:10px;padding-top:10px;font-family:'JetBrains Mono',monospace;font-size:11px;color:#4A4A45">
            {len(_asst_applied)} applied &nbsp;·&nbsp; {len(_asst_interviews)} interview{"s" if len(_asst_interviews) != 1 else ""}
          </div>
        </div>
        """, unsafe_allow_html=True)

        # Active interviews — prominent if any
        if _asst_interviews:
            st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#D4FF3A;letter-spacing:0.1em;margin-bottom:6px">ACTIVE INTERVIEWS</div>', unsafe_allow_html=True)
            for _iv in _asst_interviews:
                st.markdown(f"""
                <div style="font-family:'JetBrains Mono',monospace;font-size:10px;color:#F5F4EE;padding:5px 0;border-bottom:1px solid #1f1f22">
                  {_iv.get('title','')[:32]}<br><span style="color:#4A4A45">{_iv.get('company','')[:28]}</span>
                </div>""", unsafe_allow_html=True)
            st.markdown("<div style='margin-bottom:12px'></div>", unsafe_allow_html=True)

        # Job selector — selecting auto-primes context, no extra button
        if _asst_queue:
            st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#4A4A45;letter-spacing:0.1em;margin-bottom:6px">LOAD A JOB INTO CHAT</div>', unsafe_allow_html=True)
            _job_options = ["— select to focus chat —"] + [
                f"{j.get('score')}/10 · {j.get('title','')[:30]} @ {j.get('company','')[:18]}"
                for j in _asst_queue[:25]
            ]
            _prev_selection = st.session_state.get("_asst_prev_job_sel", "— select to focus chat —")
            _selected_label = st.selectbox(
                "job_select",
                _job_options,
                key="asst_job_selector",
                label_visibility="collapsed",
            )
            # Auto-prime on selection change — no extra button needed
            if _selected_label != "— select to focus chat —" and _selected_label != _prev_selection:
                _sel_idx = _job_options.index(_selected_label) - 1
                _sel_job = _asst_queue[_sel_idx]
                st.session_state["_asst_job_context"] = _sel_job
                st.session_state["_asst_prev_job_sel"] = _selected_label
                # Fetch company research in background
                try:
                    from researcher import load_cached, research_company as _rc
                    _cr = load_cached(_sel_job.get("company", ""))
                    if not _cr:
                        with st.spinner(f"Loading {_sel_job.get('company','')} research…"):
                            _cr = _rc(_sel_job.get("company", ""), job_title=_sel_job.get("title", ""))
                    st.session_state["_asst_company_research"] = _cr
                except Exception:
                    st.session_state["_asst_company_research"] = None
                st.rerun()

            # Show loaded job badge
            _loaded = st.session_state.get("_asst_job_context")
            if _loaded and _selected_label != "— select to focus chat —":
                st.markdown(f"""
                <div style="background:#0d1f0d;border:1px solid #1a3a1a;border-radius:6px;padding:8px 10px;margin-top:6px;font-family:'JetBrains Mono',monospace;font-size:10px">
                  <span style="color:#D4FF3A">✓ loaded</span><br>
                  <span style="color:#8B8B85">{_loaded.get('title','')[:35]}</span><br>
                  <span style="color:#4A4A45">{_loaded.get('company','')}</span>
                </div>""", unsafe_allow_html=True)
                if st.button("✕ Clear job context", key="clear_job_ctx", use_container_width=True):
                    st.session_state.pop("_asst_job_context", None)
                    st.session_state.pop("_asst_company_research", None)
                    st.session_state["_asst_prev_job_sel"] = "— select to focus chat —"
                    st.rerun()

        # Resume badge — compact, at bottom
        st.markdown(f"""
        <div style="margin-top:16px;font-family:'JetBrains Mono',monospace;font-size:10px;color:#4A4A45;padding:8px 0;border-top:1px solid #1f1f22">
          <span style="color:#D4FF3A">✓</span> Full resume in context ({len(resume_text):,} chars)
        </div>
        """, unsafe_allow_html=True)

    # ── LEFT PANEL — chat ───────────────────────────────────────────
    with chat_col:
        # Header row: title + clear button
        _h1, _h2 = st.columns([5, 1])
        with _h2:
            if st.session_state["_chat_history"]:
                if st.button("New chat", key="clear_chat", type="secondary", use_container_width=True):
                    st.session_state["_chat_history"] = []
                    if _USER_ID:
                        try:
                            clear_chat_history(_USER_ID)
                        except Exception:
                            pass
                    st.rerun()

        # Dynamic suggestions — empty state only
        if not st.session_state["_chat_history"]:
            _suggestions = _build_dynamic_suggestions(_asst_apps)
            _sc1, _sc2 = st.columns(2)
            for _i, _sug in enumerate(_suggestions):
                _sc = _sc1 if _i % 2 == 0 else _sc2
                with _sc:
                    if st.button(_sug, key=f"sug_{_i}", use_container_width=True):
                        _full = _sug
                        # Auto-inject loaded job context into suggestion if relevant
                        _jc = st.session_state.get("_asst_job_context")
                        if _jc and any(w in _sug.lower() for w in ["interview", "tailor", "application", "company"]):
                            _full = (
                                f"{_sug}\n\n[JOB CONTEXT]\nTitle: {_jc.get('title')} @ {_jc.get('company')}\n"
                                f"Score: {_jc.get('score')}/10 — {_jc.get('score_reason','')}\n"
                                f"Description:\n{(_jc.get('description') or '')[:2000]}"
                            )
                        st.session_state["_chat_history"].append({"role": "user", "content": _full})
                        _save_history()
                        st.rerun()
            st.markdown("<div style='margin-bottom:8px'></div>", unsafe_allow_html=True)

        # Render chat history
        for _msg in st.session_state["_chat_history"]:
            _avatar = "👤" if _msg["role"] == "user" else "🤖"
            # Strip injected context blocks from display — show clean user text only
            _display = re.sub(r"\n\n\[JOB CONTEXT(?:[^\]]*)?\].*", "", _msg["content"], flags=re.DOTALL).strip()
            with st.chat_message(_msg["role"], avatar=_avatar):
                st.markdown(_display)

        # Stream pending response
        _history = st.session_state["_chat_history"]
        if _history and _history[-1]["role"] == "user":
            _resp = _stream_response(_history)
            st.session_state["_chat_history"].append({"role": "assistant", "content": _resp})
            _save_history()
            st.rerun()

        # Chat input
        if _prompt := st.chat_input("Ask about your job search…"):
            _full_prompt = _prompt
            # Inject loaded job context + research into message
            _job_ctx = st.session_state.get("_asst_job_context")
            _cr_ctx  = st.session_state.get("_asst_company_research")
            if _job_ctx:
                _full_prompt = (
                    f"{_prompt}\n\n[JOB CONTEXT LOADED]\n"
                    f"Title: {_job_ctx.get('title')} @ {_job_ctx.get('company')}\n"
                    f"Score: {_job_ctx.get('score')}/10 — {_job_ctx.get('score_reason','')}\n"
                    f"Description:\n{(_job_ctx.get('description') or '')[:3000]}"
                )
                if _cr_ctx and _cr_ctx.get("research_text"):
                    _full_prompt += f"\n\n[COMPANY RESEARCH]\n{_cr_ctx['research_text']}"
            st.session_state["_chat_history"].append({"role": "user", "content": _full_prompt})
            _save_history()
            st.rerun()
